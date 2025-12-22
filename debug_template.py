
from docx import Document
import os

TEMPLATE_PATH = "tanml/report/templates/report_template_reg.docx"

if not os.path.exists(TEMPLATE_PATH):
    print("Template missing")
    exit(1)

doc = Document(TEMPLATE_PATH)

print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

print("\n--- Document Structure ---")
table_idx = 0
for i, p in enumerate(doc.paragraphs):
    print(f"P[{i}]: {p.text[:60]}")
    # Heuristic: Check if a table follows this paragraph in the internal list?
    # python-docx doesn't explicitly link p to table in linear order easily in high-level API.
    # But we can check if table is anchored here if we scan xml.
    
print("\n--- Tables ---")
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.rows[0].cells) if rows > 0 else 0
    first_cell = t.rows[0].cells[0].text if rows > 0 and cols > 0 else "EMPTY"
    print(f"Table[{i}]: {rows}x{cols}, First Cell: '{first_cell}'")
