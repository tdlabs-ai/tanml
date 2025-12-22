from docx import Document
import re
import os

template_path = "tanml/report/templates/report_template_reg.docx"
if not os.path.exists(template_path):
    print(f"Error: {template_path} not found!")
    exit(1)

doc = Document(template_path)
placeholder = "{{ModelMetaCheck.target_balance}}"
regex_pattern = r"\{\{\s*ModelMetaCheck\.target_balance\s*\}\}"

print(f"Inspecting {template_path}...")

# 1. Body Paragraphs
print("\n--- Inspecting Body Paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    if "ModelMetaCheck" in p.text or "Target Balance" in p.text:
       print(f"Found keyword in Para {i}: '{p.text}'")
       if re.search(regex_pattern, p.text):
           print(f"  -> MATCHES REGEX!")
       else:
           print(f"  -> DOES NOT MATCH REGEX. Runs: {[r.text for r in p.runs]}")

# 2. Tables
print("\n--- Inspecting Tables ---")
def traverse(tables, depth=0):
    for ti, t in enumerate(tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                # Check directly
                if "ModelMetaCheck" in cell.text or "Target Balance" in cell.text:
                    print(f"Found keyword in Table (depth {depth}) T{ti} R{ri} C{ci}: '{cell.text.strip()}'")
                    # Check paragraphs in cell
                    for pi, p in enumerate(cell.paragraphs):
                        if "ModelMetaCheck" in p.text:
                             print(f"  -> In Para {pi}: '{p.text}'")
                             if re.search(regex_pattern, p.text):
                                 print(f"    -> MATCHES REGEX!")
                             else:
                                 print(f"    -> DOES NOT MATCH REGEX. Runs: {[r.text for r in p.runs]}")
                
                if cell.tables:
                    traverse(cell.tables, depth=depth+1)

traverse(doc.tables)
