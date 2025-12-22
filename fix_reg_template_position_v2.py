
from docx import Document
import os

TEMPLATE_PATH = "tanml/report/templates/report_template_reg.docx"

def fix_position(doc):
    # 1. CLEANUP (Remove old tables/headers at END or anywhere)
    to_delete_p = []
    tables_to_delete = []
    
    # Header Cleanup
    for p in doc.paragraphs:
        if "2a. Cross-Validation" in p.text:
            print(f"Deleting misplaced header: {p.text}")
            to_delete_p.append(p)
            
    # Table Cleanup
    for tbl in doc.tables:
        try:
            if tbl.rows[0].cells[0].text.strip() == "Metric":
                if "RMSE" in [r.cells[0].text for r in tbl.rows[1:]]:
                    print("Deleting misplaced CV Table")
                    tables_to_delete.append(tbl)
        except:
            pass
            
    for t in tables_to_delete:
        t._element.getparent().remove(t._element)
    for p in to_delete_p:
        p._element.getparent().remove(p._element)

    # 2. INSERT
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Performance Metrics" in p.text:
            print(f"Found target anchor at index {i}: '{p.text}'")
            target_idx = i
            break
            
    if target_idx == -1:
        print("Error: Anchor not found!")
        return

    # specific insertion point: Before the NEXT paragraph
    # If target is P[12], we want to insert before P[13].
    
    if target_idx + 1 < len(doc.paragraphs):
        anchor_p = doc.paragraphs[target_idx + 1]
        print(f"Inserting BEFORE: '{anchor_p.text}' (Index {target_idx+1})")
        
        # Insert Header (Manual styling to match 'Normal' bold)
        h = anchor_p.insert_paragraph_before("2a. Cross-Validation Performance Summary (Robust)")
        run = h.runs[0] if h.runs else h.add_run("2a. Cross-Validation Performance Summary (Robust)")
        if not h.runs: run.text = "2a. Cross-Validation Performance Summary (Robust)"
        run.bold = True
        
        # Insert Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Move table to be after h (and thus before anchor_p)
        # Standard add_table adds to end. We move it.
        h._p.addnext(table._tbl)
        
        # Fill Table
        hdr = table.rows[0].cells
        cols = ["Metric", "Mean", "Std Dev", "P05", "P50", "P95"]
        for j, c in enumerate(cols):
            hdr[j].text = c
            try: hdr[j].paragraphs[0].runs[0].font.bold = True
            except: pass
            
        metrics = [("RMSE", "rmse"), ("MAE", "mae"), ("Median AE", "median_ae"), ("R2 Score", "r2")]
        for label, key in metrics:
            row = table.add_row().cells
            row[0].text = label
            try: row[0].paragraphs[0].runs[0].font.bold = True
            except: pass
            row[1].text = f"{{{{cv.{key}.mean}}}}"
            row[2].text = f"{{{{cv.{key}.std}}}}"
            row[3].text = f"{{{{cv.{key}.p05}}}}"
            row[4].text = f"{{{{cv.{key}.p50}}}}"
            row[5].text = f"{{{{cv.{key}.p95}}}}"

        print("Insertion complete.")
        
    else:
        print("Target was last paragraph. Appending.")
        # ... logic for append if needed, but unlikely given debug output ...

if __name__ == "__main__":
    doc = Document(TEMPLATE_PATH)
    fix_position(doc)
    doc.save(TEMPLATE_PATH)
