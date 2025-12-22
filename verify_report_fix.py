from docx import Document
import os

def test_insert_after():
    # Create valid docx
    doc = Document()
    doc.add_paragraph("Header")
    p = doc.add_paragraph("Here is the anchor text for plots.")
    doc.add_paragraph("Footer")
    
    anchor = "anchor"
    
    print("Testing insert logic...")
    found = False
    for p in doc.paragraphs:
        if anchor.lower() in p.text.lower():
            found = True
            # The fixed logic:
            new_p = p._parent.add_paragraph()
            new_p.add_run("(no plots available)")
            p._p.addnext(new_p._p)
            break
            
    if not found:
        print("❌ Anchor not found in test doc!")
        return

    # Verify order
    texts = [p.text for p in doc.paragraphs]
    print("Paragraphs:", texts)
    
    expected_index = texts.index("Here is the anchor text for plots.") + 1
    if texts[expected_index] == "(no plots available)":
        print("✅ Success: Paragraph inserted correctly after anchor.")
    else:
        print(f"❌ Failed: Expected '(no plots available)' at index {expected_index}, got '{texts[expected_index] if len(texts) > expected_index else 'None'}'")

if __name__ == "__main__":
    test_insert_after()
