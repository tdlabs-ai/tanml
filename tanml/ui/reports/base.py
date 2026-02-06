# tanml/ui/reports/base.py
"""
Base classes for modular report generation.

This module provides the foundation for creating extensible report sections.
Contributors can add new report sections by implementing the ReportSection protocol.

Example:
    from tanml.ui.reports.base import ReportSection

    class MyAnalysisSection(ReportSection):
        name = "My Analysis"
        order = 50  # Controls position in report

        def should_include(self, context: ReportContext) -> bool:
            return "my_analysis" in context.data

        def add_to_document(self, doc: Document, context: ReportContext) -> None:
            doc.add_heading(self.name, level=2)
            doc.add_paragraph(context.data["my_analysis"]["summary"])
"""

from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx import Document


@dataclass
class ReportContext:
    """
    Context object passed to report sections.

    Contains all data needed for report generation, including
    model info, metrics, plots, and analysis results.
    """

    task_type: str
    model_name: str = "Model"
    data: dict[str, Any] = field(default_factory=dict)
    plots: dict[str, Any] = field(default_factory=dict)
    timestamp: str = ""

    def get(self, key: str, default: Any = None) -> Any:
        """Get data from context."""
        return self.data.get(key, default)

    def has(self, key: str) -> bool:
        """Check if data exists in context."""
        return key in self.data


class ReportSection(ABC):
    """
    Abstract base class for report sections.

    Each section is responsible for adding its content to a Word document.
    Sections are ordered by their `order` attribute.
    """

    name: str = "Section"
    order: int = 100

    @abstractmethod
    def add_to_document(self, doc: Document, context: ReportContext) -> None:
        """Add this section's content to the Word document."""
        pass

    def should_include(self, context: ReportContext) -> bool:
        """Determine whether this section should be included."""
        return True


class SectionRegistry:
    """
    Registry for report sections.

    Allows dynamic registration and discovery of report sections.
    """

    _sections: dict[str, dict[str, type]] = {
        "development": {},
        "evaluation": {},
        "ranking": {},
    }

    @classmethod
    def register(cls, report_type: str):
        """Decorator to register a section for a report type."""

        def decorator(section_class: type) -> type:
            cls._sections.setdefault(report_type, {})[section_class.name] = section_class
            return section_class

        return decorator

    @classmethod
    def get_for_report(cls, report_type: str) -> list[type]:
        """Get all sections for a report type, sorted by order."""
        sections = cls._sections.get(report_type, {})
        return sorted(sections.values(), key=lambda s: s.order)


def add_table_with_borders(doc: Document, data: list[list[str]], headers: bool = True):
    """Add a table with borders to the document."""
    if not data:
        return

    rows = len(data)
    cols = len(data[0])

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)

            if headers and i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_image_from_figure(doc: Document, fig, width_inches: float = 5.0):
    """Add a matplotlib figure to the document."""
    import io

    from docx.shared import Inches

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)

    doc.add_picture(buf, width=Inches(width_inches))
    buf.close()
