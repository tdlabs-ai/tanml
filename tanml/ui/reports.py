# tanml/ui/reports.py
"""
Report generation functions for TanML UI.

These functions generate Word document reports from validation results.
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.colors as mcolors

from tanml.ui.glossary import GLOSSARY
from tanml.ui.narratives import (
    story_performance,
    story_features,
    story_overfitting,
    story_drift,
    story_stress,
    story_shap,
)


def generate_dev_report_docx(dev_data: Dict[str, Any]) -> io.BytesIO:
    """
    Generate a Model Development Report (Word document).
    
    Args:
        dev_data: Dictionary containing development metrics and plots
        
    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()
    doc.add_heading("Model Development Report", 0)
    
    # Model Info Section
    model_info = dev_data.get("model_info", {})
    doc.add_heading("1. Model Information", level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Property"
    hdr_cells[1].text = "Value"
    
    for key, value in model_info.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    # Cross-Validation Results
    cv_results = dev_data.get("cv_results", {})
    if cv_results:
        doc.add_heading("2. Cross-Validation Results", level=1)
        
        cv_summary = cv_results.get("summary", {})
        if cv_summary:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Mean"
            hdr_cells[2].text = "Std"
            hdr_cells[3].text = "Range (P05-P95)"
            
            for metric, stats in cv_summary.items():
                if isinstance(stats, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = metric.upper()
                    row_cells[1].text = f"{stats.get('mean', 0):.4f}"
                    row_cells[2].text = f"{stats.get('std', 0):.4f}"
                    p05 = stats.get('p05', 0)
                    p95 = stats.get('p95', 0)
                    row_cells[3].text = f"{p05:.4f} - {p95:.4f}"
    
    # Feature Importance
    feature_importance = dev_data.get("feature_importance", [])
    if feature_importance:
        doc.add_heading("3. Feature Importance", level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Feature"
        hdr_cells[1].text = "Importance"
        
        for feat in feature_importance[:20]:  # Top 20
            row_cells = table.add_row().cells
            row_cells[0].text = str(feat.get("feature", ""))
            row_cells[1].text = f"{feat.get('importance', 0):.4f}"
    
    # Add plots
    plots = dev_data.get("plots", {})
    if plots:
        doc.add_heading("4. Diagnostic Plots", level=1)
        
        for plot_name, plot_path in plots.items():
            if plot_path and os.path.exists(plot_path):
                doc.add_paragraph(plot_name.replace("_", " ").title())
                doc.add_picture(plot_path, width=Inches(6))
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_eval_report_docx(buf: Dict[str, Any]) -> io.BytesIO:
    """
    Generate a Model Evaluation Report (Word document).
    
    Args:
        buf: Dictionary containing evaluation metrics and plots
        
    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()
    doc.add_heading("Model Evaluation Report", 0)
    
    # Glossary Section
    doc.add_heading("Glossary", level=1)
    for term, definition in list(GLOSSARY.items())[:10]:  # First 10 terms
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)
    
    # Performance Comparison
    doc.add_heading("1. Performance Comparison", level=1)
    
    train_metrics = buf.get("train_metrics", {})
    test_metrics = buf.get("test_metrics", {})
    task_type = buf.get("task_type", "classification")
    
    # Narrative summary
    if test_metrics:
        narrative = story_performance(test_metrics, task_type)
        doc.add_paragraph(narrative)
    
    # Metrics table
    if train_metrics or test_metrics:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Train"
        hdr_cells[2].text = "Test"
        
        all_metrics = set(train_metrics.keys()) | set(test_metrics.keys())
        for metric in sorted(all_metrics):
            row_cells = table.add_row().cells
            row_cells[0].text = metric.upper()
            train_val = train_metrics.get(metric)
            test_val = test_metrics.get(metric)
            row_cells[1].text = f"{train_val:.4f}" if train_val is not None else "—"
            row_cells[2].text = f"{test_val:.4f}" if test_val is not None else "—"
        
        # Overfitting analysis
        if train_metrics and test_metrics:
            narrative = story_overfitting(train_metrics, test_metrics)
            doc.add_paragraph(narrative)
    
    # Diagnostic Plots
    plots = buf.get("plots", {})
    if plots:
        doc.add_heading("2. Diagnostic Plots", level=1)
        
        plot_pairs = [
            ("roc_train", "roc_test"),
            ("pr_train", "pr_test"),
            ("confusion_train", "confusion_test"),
        ]
        
        for train_key, test_key in plot_pairs:
            train_path = plots.get(train_key)
            test_path = plots.get(test_key)
            
            # Add side by side if both exist
            if train_path and os.path.exists(train_path):
                doc.add_paragraph(f"{train_key.replace('_', ' ').title()}")
                doc.add_picture(train_path, width=Inches(3))
            
            if test_path and os.path.exists(test_path):
                doc.add_picture(test_path, width=Inches(3))
    
    # Risk Assessment
    doc.add_heading("3. Risk Assessment", level=1)
    
    # Drift analysis
    drift_data = buf.get("drift_analysis", {})
    drift_narrative = story_drift(drift_data)
    doc.add_paragraph(drift_narrative)
    
    # Stress test
    stress_data = buf.get("stress_test", [])
    stress_narrative = story_stress(stress_data)
    doc.add_paragraph(stress_narrative)
    
    # Explainability
    doc.add_heading("4. Model Explainability", level=1)
    
    shap_results = buf.get("shap_results", {})
    shap_narrative = story_shap(shap_results)
    doc.add_paragraph(shap_narrative)
    
    # SHAP plot
    shap_plot = plots.get("shap_beeswarm")
    if shap_plot and os.path.exists(shap_plot):
        doc.add_picture(shap_plot, width=Inches(6))
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_ranking_report_docx(
    metrics_df: pd.DataFrame,
    corr_df: pd.DataFrame,
    method: str,
    target: str,
    task_type: str,
    X: Optional[pd.DataFrame] = None,
    y: Optional[pd.Series] = None,
) -> io.BytesIO:
    """
    Generate a Feature Ranking Report (Word document).
    
    Args:
        metrics_df: DataFrame with feature ranking scores
        corr_df: Correlation matrix DataFrame
        method: Ranking method used
        target: Target column name
        task_type: "classification" or "regression"
        X: Optional features DataFrame
        y: Optional target Series
        
    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()
    doc.add_heading("Feature Ranking Report", 0)
    
    # Summary
    doc.add_heading("1. Summary", level=1)
    doc.add_paragraph(f"**Method**: {method}")
    doc.add_paragraph(f"**Target**: {target}")
    doc.add_paragraph(f"**Task Type**: {task_type.title()}")
    if X is not None:
        doc.add_paragraph(f"**Features**: {len(X.columns)}")
        doc.add_paragraph(f"**Samples**: {len(X)}")
    
    # Feature Rankings
    if metrics_df is not None and not metrics_df.empty:
        doc.add_heading("2. Feature Rankings", level=1)
        
        # Narrative
        narrative = story_features(metrics_df, top_n=5)
        doc.add_paragraph(narrative)
        
        # Table (top 30)
        display_df = metrics_df.head(30)
        
        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = "Table Grid"
        
        # Headers
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(display_df.columns):
            hdr_cells[i].text = str(col)
        
        # Data rows
        for _, row in display_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
    
    # Correlation Analysis
    if corr_df is not None and not corr_df.empty:
        doc.add_heading("3. Feature Correlations", level=1)
        
        # Find high correlations
        high_corr_pairs = []
        for i in range(len(corr_df.columns)):
            for j in range(i + 1, len(corr_df.columns)):
                corr = corr_df.iloc[i, j]
                if abs(corr) > 0.7:
                    high_corr_pairs.append((
                        corr_df.columns[i],
                        corr_df.columns[j],
                        corr
                    ))
        
        if high_corr_pairs:
            doc.add_paragraph(f"Found {len(high_corr_pairs)} highly correlated pairs (|r| > 0.7):")
            
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Feature 1"
            hdr_cells[1].text = "Feature 2"
            hdr_cells[2].text = "Correlation"
            
            for f1, f2, corr in sorted(high_corr_pairs, key=lambda x: -abs(x[2]))[:20]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(f1)
                row_cells[1].text = str(f2)
                row_cells[2].text = f"{corr:.4f}"
        else:
            doc.add_paragraph("No highly correlated feature pairs found (|r| > 0.7).")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Legacy aliases
_generate_dev_report_docx = generate_dev_report_docx
_generate_eval_report_docx = generate_eval_report_docx
_generate_ranking_report_docx = generate_ranking_report_docx
