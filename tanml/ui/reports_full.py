# tanml/ui/reports_full.py
"""
Full report generation functions for TanML UI.

These functions generate Word document reports from validation results.
This module is imported by app.py to reduce its size.
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import seaborn as sns

from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from tanml.ui.glossary import GLOSSARY
from tanml.ui.narratives import (
    story_performance as _story_performance,
    story_features as _story_features,
    story_overfitting as _story_overfitting,
    story_drift as _story_drift,
    story_stress as _story_stress,
    story_shap as _story_shap,
)


def generate_dev_report_docx(dev_data: Dict[str, Any]) -> io.BytesIO:
    """
    Generate a Model Development Report DOCX.
    
    Args:
        dev_data: Development data containing metrics, images, cv_metrics, etc.
        
    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()
    doc.add_heading('Model Development Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Glossary / Guide
    doc.add_heading('Guide to Metrics', level=2)
    doc.add_paragraph("Key terms used in this report:")
    
    # 1. General Concepts
    if dev_data.get("cv_metrics"): 
        doc.add_paragraph(f"**Cross-Validation**: {GLOSSARY.get('Cross-Validation', '')}")
    
    # 2. Metrics relevant to task
    metrics_to_show = []
    if dev_data.get("task_type") == "classification":
        metrics_to_show = ["ROC AUC", "PR AUC", "F1 Score", "Accuracy", "Precision", "Recall", "Log Loss", "Brier Score", "MCC", "KS Statistic", "Confusion Matrix"]
    else:
        metrics_to_show = ["RMSE", "MAE", "R2 Score", "Median AE"]
        
    for m in metrics_to_show:
        if m in GLOSSARY:
             doc.add_paragraph(f"**{m}**: {GLOSSARY[m]}", style='List Bullet')
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph()
    if dev_data:
        task = dev_data.get("task_type", "Unknown")
        algo = dev_data.get("model_config", {}).get("algorithm", "Unknown")
        p.add_run(f"Summarizes development of a ").bold = False
        p.add_run(f"{task.title()}").bold = True
        p.add_run(f" model using ").bold = False
        p.add_run(f"{algo}").bold = True
        p.add_run(".")
    else:
        p.add_run("No development data available.")
        
    # 2. Results
    if dev_data:
        doc.add_heading('2. Model Performance', level=1)
        
        # Narrative
        narrative = _story_performance(dev_data.get("metrics", {}), dev_data.get("task_type"))
        if narrative:
            doc.add_heading('Executive Insight', level=2)
            doc.add_paragraph(narrative)
            
        doc.add_heading('Final Model Metrics (Full Dataset)', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Metric'; hdr[1].text = 'Score'
        for k, v in dev_data.get("metrics", {}).items():
            r = table.add_row().cells
            r[0].text = k.replace("_", " ").title()
            r[1].text = f"{v:.4f}" if isinstance(v, (float, int)) else str(v)
            
        # CV Metrics
        cv_met = dev_data.get("cv_metrics", {})
        if cv_met:
            doc.add_heading('Cross-Validation Metrics (Mean)', level=2)
            table2 = doc.add_table(rows=1, cols=2)
            table2.style = 'Table Grid'
            hdr2 = table2.rows[0].cells
            hdr2[0].text = 'Metric'; hdr2[1].text = 'Mean Score'
            for k, v in cv_met.items():
                r = table2.add_row().cells
                r[0].text = k.replace("_", " ").title()
                r[1].text = f"{v:.4f}" if isinstance(v, (float, int)) else str(v)

        # Images
        imgs = dev_data.get("images", {})
        cv_imgs = dev_data.get("cv_images", {})
        
        if cv_imgs:
            doc.add_heading('2.1 Cross-Validation Plots', level=2)
            for name, img_bytes in cv_imgs.items():
                    doc.add_paragraph(name.replace("_", " ").title())
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))
                    
        if imgs:
            doc.add_heading('2.2 Final Model Diagnostics', level=2)
            for name, img_bytes in imgs.items():
                    doc.add_paragraph(name.replace("_", " ").title())
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_eval_report_docx(buf: Dict[str, Any]) -> io.BytesIO:
    """
    Generate a Model Evaluation Report DOCX.
    
    Args:
        buf: Evaluation data buffer containing evaluation, drift, stress, etc.
        
    Returns:
        BytesIO buffer containing the Word document
    """
    # Handle case where buf might be a list or non-dict
    if isinstance(buf, list):
        buf = buf[0] if buf else {}
    if not isinstance(buf, dict):
        buf = {}
    
    doc = Document()
    doc.add_heading('Model Evaluation Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Glossary
    doc.add_heading('Guide to Metrics', level=2)
    doc.add_paragraph("Key terms used in this report:")

    
    # Determine task type from evaluation data
    ev = buf.get("evaluation", {})
    # Handle case where ev might be a list or non-dict
    if isinstance(ev, list):
        ev = ev[0] if ev else {}
    if not isinstance(ev, dict):
        ev = {}
    task_type = ev.get("task_type", "classification") if ev else "classification"

    
    # Performance Metrics
    if task_type == "classification":
        cls_metrics = ["ROC AUC", "PR AUC", "F1 Score", "Accuracy", "Precision", "Recall", "Log Loss", "Brier Score", "MCC", "KS Statistic"]
        for m in cls_metrics:
            if m in GLOSSARY:
                doc.add_paragraph(f"**{m}**: {GLOSSARY[m]}", style='List Bullet')
    else:
        reg_metrics = ["RMSE", "MAE", "R2 Score", "Median AE"]
        for m in reg_metrics:
            if m in GLOSSARY:
                doc.add_paragraph(f"**{m}**: {GLOSSARY[m]}", style='List Bullet')
    
    # Validation Concepts (always included)
    for key in ["PSI", "Stress Test", "SHAP", "Cluster Coverage"]:
        if key in GLOSSARY:
            doc.add_paragraph(f"**{key}**: {GLOSSARY[key]}", style='List Bullet')
    
    # 1. Comparison
    if ev:
        doc.add_heading('1. Results Comparison (Train vs Test)', level=1)
        
        # Narrative
        m_tr = ev.get("metrics_train", {})
        m_te = ev.get("metrics_test", {})
        narrative = _story_overfitting(m_tr, m_te)
        if narrative:
            doc.add_heading('Stability Check', level=2)
            doc.add_paragraph(narrative)
        
        # Metrics Table
        doc.add_heading('Metrics', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        h = table.rows[0].cells
        h[0].text = 'Metric'; h[1].text = 'Train'; h[2].text = 'Test'
        
        all_keys = set(m_tr.keys()) | set(m_te.keys())
        for k in sorted(all_keys):
            if k in ["confusion_matrix", "curves", "threshold_info"]: continue
            r = table.add_row().cells
            r[0].text = k
            r[1].text = f"{m_tr.get(k, 0):.4f}"
            r[2].text = f"{m_te.get(k, 0):.4f}"
        
        # Diagnostic Plots
        ev_imgs = ev.get("images", {})
        if ev_imgs:
            doc.add_heading('Diagnostic Plots', level=2)
            _add_eval_plots(doc, ev_imgs)
    else:
        doc.add_paragraph("No evaluation data found.")

    # 2. Risk (Drift/Stress)
    doc.add_heading('2. Risk Assessment', level=1)
    
    # Drift
    drift_data = buf.get("drift")
    drift_imgs = buf.get("drift_images", {})
    doc.add_heading('2.1 Drift Analysis', level=2)
    
    nar = _story_drift(drift_data)
    if nar:
        doc.add_paragraph(nar)
    
    if drift_data:
        _add_drift_table(doc, drift_data)
        if "top_distribution" in drift_imgs:
            doc.add_paragraph("Top Drifting Feature Distribution")
            doc.add_picture(io.BytesIO(drift_imgs["top_distribution"]), width=Inches(5))
    else:
        doc.add_paragraph("Drift analysis not run.")
    
    # Stress
    stress_data = buf.get("stress")
    # Handle case where stress_data might not be a list of dicts
    if stress_data and not isinstance(stress_data, list):
        stress_data = [stress_data] if isinstance(stress_data, dict) else None
    doc.add_heading('2.2 Stress Testing', level=2)

    
    nar_stress = _story_stress(stress_data)
    if nar_stress:
        doc.add_paragraph(nar_stress)
        
    if stress_data:
        _add_stress_table(doc, stress_data)
    else:
        doc.add_paragraph("Stress test not run.")

    # Cluster Coverage
    cluster_data = buf.get("cluster_coverage")
    cluster_imgs = buf.get("cluster_images", {})
    doc.add_heading('2.3 Input Cluster Coverage Check', level=2)
    
    if cluster_data:
        _add_cluster_section(doc, cluster_data, cluster_imgs)
    else:
        doc.add_paragraph("Cluster coverage check not run.")

    # Benchmarking
    bench_data = buf.get("benchmark")
    bench_imgs = buf.get("benchmark_images", {})
    doc.add_heading('2.4 Benchmarking', level=2)
    
    if bench_data:
        _add_benchmark_section(doc, bench_data, bench_imgs)
    else:
        doc.add_paragraph("Benchmarking not run.")

    # 3. Explainability
    doc.add_heading('3. Explainability', level=1)
    expl = buf.get("explainability", {})
    # Handle case where expl might be a list or non-dict
    if isinstance(expl, list):
        expl = expl[0] if expl else {}
    if not isinstance(expl, dict):
        expl = {}
    if expl and expl.get("status") == "ok":
        _add_explainability_section(doc, expl)
    else:
        doc.add_paragraph("Explainability not run.")


    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_eval_plots(doc: Document, ev_imgs: Dict[str, bytes]) -> None:
    """Add evaluation plots with train/test pairs."""
    unique_bases = set()
    for name in ev_imgs.keys():
        if name.startswith("train_"): unique_bases.add(name.replace("train_", ""))
        elif name.startswith("test_"): unique_bases.add(name.replace("test_", ""))
        else: unique_bases.add(name)
    
    for base in sorted(unique_bases):
        train_key = f"train_{base}"
        test_key = f"test_{base}"
        
        if train_key in ev_imgs and test_key in ev_imgs:
            doc.add_heading(base.replace("_", " ").title(), level=3)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            r = table.rows[0].cells
            
            p1 = r[0].add_paragraph("Train")
            p1.alignment = 1
            r[0].add_paragraph().add_run().add_picture(io.BytesIO(ev_imgs[train_key]), width=Inches(3.0))
            
            p2 = r[1].add_paragraph("Test")
            p2.alignment = 1
            r[1].add_paragraph().add_run().add_picture(io.BytesIO(ev_imgs[test_key]), width=Inches(3.0))
        
        elif base in ev_imgs:
            doc.add_paragraph(base.replace("_", " ").title())
            doc.add_picture(io.BytesIO(ev_imgs[base]), width=Inches(4))


def _add_drift_table(doc: Document, drift_data) -> None:
    """Add drift analysis table."""
    # Defensive null/type check
    if not drift_data:
        doc.add_paragraph("No drift data available.")
        return
    if not isinstance(drift_data, list):
        drift_data = [drift_data] if isinstance(drift_data, dict) else []
    if not drift_data:
        doc.add_paragraph("No drift data available.")
        return
        
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    h = table.rows[0].cells
    h[0].text = 'Feature'; h[1].text = 'PSI'; h[2].text = 'PSI Status'; h[3].text = 'KS Stat'; h[4].text = 'KS Status'
    for d in drift_data:
        if not isinstance(d, dict):
            continue
        r = table.add_row().cells
        r[0].text = str(d.get('Feature', ''))
        r[1].text = f"{d.get('PSI', 0):.4f}"
        r[2].text = str(d.get('PSI Status', d.get('Status', ''))).replace("🟢 ", "").replace("🟠 ", "").replace("🔴 ", "").replace("🟡 ", "").replace("● ", "")
        r[3].text = f"{d.get('KS Stat', 0):.4f}"
        r[4].text = str(d.get('KS Status', '')).replace("🟢 ", "").replace("🟠 ", "").replace("🔴 ", "").replace("🟡 ", "")


def _add_stress_table(doc: Document, stress_data) -> None:
    """Add stress test results table."""
    # Defensive null/type check
    if not stress_data:
        doc.add_paragraph("No stress test data available.")
        return
    if not isinstance(stress_data, list):
        stress_data = [stress_data] if isinstance(stress_data, dict) else []
    if not stress_data or not stress_data[0]:
        doc.add_paragraph("No stress test data available.")
        return
    
    first_row = stress_data[0]
    if not isinstance(first_row, dict):
        doc.add_paragraph("Invalid stress test data format.")
        return
        
    doc.add_paragraph("Stress Test Results (Perturbed Data)")
    table = doc.add_table(rows=1, cols=len(first_row))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, k in enumerate(first_row.keys()): hdr[i].text = str(k)
    for row_dat in stress_data:
        if not isinstance(row_dat, dict):
            continue
        r = table.add_row().cells
        for i, v in enumerate(row_dat.values()): r[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)


def _add_cluster_section(doc: Document, cluster_data: Dict, cluster_imgs: Dict) -> None:
    """Add cluster coverage section."""
    coverage_pct = cluster_data.get("coverage_pct", 0)
    ood_pct = cluster_data.get("ood_pct", 0)
    n_clusters = cluster_data.get("n_clusters", 0)
    uncovered = cluster_data.get("uncovered_clusters", 0)
    
    if coverage_pct >= 95:
        doc.add_paragraph(f"✅ **Excellent Coverage**: Test data covers {coverage_pct:.1f}% of the {n_clusters} training input space clusters.")
    elif coverage_pct >= 80:
        doc.add_paragraph(f"⚠️ **Good Coverage**: Test data covers {coverage_pct:.1f}% of clusters ({uncovered} uncovered).")
    else:
        doc.add_paragraph(f"🚨 **Poor Coverage**: Only {coverage_pct:.1f}% of training clusters are covered by test data.")
    
    if ood_pct > 10:
        doc.add_paragraph(f"⚠️ **OOD Alert**: {ood_pct:.1f}% of test samples appear to be out-of-distribution.")
    
    # Summary table
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ("Total Clusters", str(n_clusters)),
        ("Coverage", f"{coverage_pct:.1f}%"),
        ("Covered Clusters", str(cluster_data.get("covered_clusters", 0))),
        ("Uncovered Clusters", str(uncovered)),
        ("OOD Samples", f"{ood_pct:.1f}% ({cluster_data.get('ood_count', 0)} samples)")
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    # Images
    if "distribution" in cluster_imgs:
        doc.add_paragraph("Cluster Distribution Chart:")
        doc.add_picture(io.BytesIO(cluster_imgs["distribution"]), width=Inches(5))
    
    if "pca_scatter" in cluster_imgs:
        doc.add_paragraph("Cluster Space Visualization (PCA):")
        doc.add_picture(io.BytesIO(cluster_imgs["pca_scatter"]), width=Inches(5))


def _add_benchmark_section(doc: Document, bench_data: Dict, bench_imgs: Dict) -> None:
    """Add benchmarking section."""
    your_model = bench_data.get("your_model", {})
    baselines = bench_data.get("baselines", {})
    
    if baselines:
        doc.add_paragraph(f"Comparison of your model against {len(baselines)} baseline model(s) on test data.")
        
        metrics = list(your_model.keys())
        n_cols = 2 + len(baselines)
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Your Model"
        for i, baseline_name in enumerate(baselines.keys()):
            hdr[2 + i].text = baseline_name[:20]
        
        for metric in metrics:
            r = table.add_row().cells
            r[0].text = metric.upper()
            r[1].text = f"{your_model.get(metric, 0):.4f}"
            for i, baseline_name in enumerate(baselines.keys()):
                r[2 + i].text = f"{baselines[baseline_name].get(metric, 0):.4f}"
        
        if "comparison" in bench_imgs:
            doc.add_paragraph("Performance Comparison Chart:")
            doc.add_picture(io.BytesIO(bench_imgs["comparison"]), width=Inches(6))
    else:
        doc.add_paragraph("No baseline models were compared.")


def _add_explainability_section(doc: Document, expl: Dict) -> None:
    """Add explainability section."""
    nar_shap = _story_shap(expl)
    if nar_shap:
        doc.add_heading('Executive Insight', level=2)
        doc.add_paragraph(nar_shap)
        
    plots = expl.get("plots", {})
    
    if "beeswarm" in plots:
        doc.add_heading('SHAP Beeswarm', level=2)
        p = plots["beeswarm"]
        if isinstance(p, str) and os.path.exists(p): 
            doc.add_picture(p, width=Inches(5))
        elif isinstance(p, bytes): 
            doc.add_picture(io.BytesIO(p), width=Inches(5))
            
    if "bar" in plots:
        doc.add_heading('SHAP Importance', level=2)
        p = plots["bar"]
        if isinstance(p, str) and os.path.exists(p): 
            doc.add_picture(p, width=Inches(5))
        elif isinstance(p, bytes): 
            doc.add_picture(io.BytesIO(p), width=Inches(5))
        
    if "top_features" in expl:
        doc.add_heading('Top Drivers Table', level=2)
        tf = expl["top_features"]
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Feature"
        table.rows[0].cells[1].text = "Impact"
        for row in tf[:5]:
            r = table.add_row().cells
            k_feat = "feature" if "feature" in row else list(row.keys())[0]
            k_imp = "importance" if "importance" in row else list(row.keys())[1]
            r[0].text = str(row.get(k_feat, ""))
            val = row.get(k_imp, 0)
            r[1].text = f"{val:.4f}" if isinstance(val, float) else str(val)


def generate_ranking_report_docx(
    metrics_df: pd.DataFrame,
    corr_df: pd.DataFrame,
    method: str,
    target: str,
    task_type: str,
    X: Optional[pd.DataFrame] = None,
    y: Optional[pd.Series] = None,
) -> io.BytesIO:
    """
    Generate a Feature Power Ranking Report DOCX.
    
    Args:
        metrics_df: Feature importance metrics DataFrame
        corr_df: Correlation matrix DataFrame
        method: Ranking method used
        target: Target column name
        task_type: "classification" or "regression"
        X: Optional features DataFrame for plots
        y: Optional target Series for plots
        
    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()
    doc.add_heading('Feature Power Ranking Report', 0)
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(f"Target Variable: {target}")
    doc.add_paragraph(f"Task Type: {task_type}")
    doc.add_paragraph(f"Ranking Method: {method}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading('Metric Definition', level=2)
    if "Power Score" in GLOSSARY:
        doc.add_paragraph(f"**Power Score**: {GLOSSARY['Power Score']}", style='List Bullet')
    
    # Automated Narrative
    narrative = _story_features(metrics_df)
    if narrative:
        doc.add_heading('Key Insights', level=2)
        doc.add_paragraph(narrative)
    
    doc.add_heading('Top Influential Features:', level=2)
    if metrics_df is not None and not metrics_df.empty:
        top_5 = metrics_df.head(5)
        for _, row in top_5.iterrows():
            doc.add_paragraph(f"{row['Feature']}: Power Score {row['Power']:.1f}", style='List Bullet')

        # 2. Detailed Metrics
        doc.add_heading('2. Feature Importance Metrics', level=1)
        
        t = doc.add_table(rows=1, cols=len(metrics_df.columns))
        t.style = 'Table Grid'
        
        hdr_cells = t.rows[0].cells
        for i, col_name in enumerate(metrics_df.columns):
            hdr_cells[i].text = str(col_name)
            
        for _, row in metrics_df.iterrows():
            row_cells = t.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.3f}"
                else:
                    row_cells[i].text = str(val)

    # 3. Visuals
    if X is not None and y is not None and metrics_df is not None:
        doc.add_heading('3. Key Feature Distributions', level=1)
        doc.add_paragraph(f"Distribution overlay for top features relative to target '{target}'.")
        
        sns.set_theme(style="whitegrid")
        
        for idx, row in metrics_df.head(5).iterrows():
            feat = row['Feature']
            if feat not in X.columns: continue
            
            try:
                fig, ax = plt.subplots(figsize=(6, 4))
                
                if task_type == 'classification':
                    sns.kdeplot(data=X, x=feat, hue=y, fill=True, common_norm=False, palette="tab10", ax=ax)
                    ax.set_title(f"{feat} distribution by {target}")
                else:
                    sns.scatterplot(x=X[feat], y=y, alpha=0.6, ax=ax)
                    ax.set_title(f"{feat} vs {target}")
                    ax.set_ylabel(target)
                
                plt.tight_layout()
                
                img_buf = io.BytesIO()
                fig.savefig(img_buf, format='png', dpi=100)
                plt.close(fig)
                img_buf.seek(0)
                
                doc.add_heading(f"Feature: {feat}", level=3)
                doc.add_picture(img_buf, width=Inches(5.0))
                
            except Exception as e:
                doc.add_paragraph(f"Could not plot {feat}: {e}")

    # 4. Correlation Matrix
    if corr_df is not None and not corr_df.empty:
        doc.add_heading('4. Correlation Matrix', level=1)
        doc.add_paragraph("Pairwise correlation of numeric features (Pearson).")
        
        c_disp = corr_df.reset_index().rename(columns={"index": "Feature"})
        
        cmap = plt.get_cmap("coolwarm")
        d_min = corr_df.min().min()
        d_max = corr_df.max().max()
        norm = mcolors.Normalize(vmin=d_min, vmax=d_max)
        
        t2 = doc.add_table(rows=1, cols=len(c_disp.columns))
        t2.style = 'Table Grid'
        
        hdr2 = t2.rows[0].cells
        for i, col_name in enumerate(c_disp.columns):
            hdr2[i].text = str(col_name)
            
        for _, row in c_disp.iterrows():
            row_cells = t2.add_row().cells
            for i, val in enumerate(row):
                if i == 0:
                    row_cells[i].text = str(val)
                    continue

                if isinstance(val, (int, float)):
                    row_cells[i].text = f"{val:.2f}"
                    
                    try:
                        rgba = cmap(norm(val))
                        hex_color = mcolors.to_hex(rgba, keep_alpha=False).lstrip('#')
                        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
                        row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

                        r, g, b, _ = rgba
                        lum = 0.299*r + 0.587*g + 0.114*b
                        if lum < 0.5:
                           run = row_cells[i].paragraphs[0].runs[0]
                           run.font.color.rgb = RGBColor(255, 255, 255)
                           
                    except Exception:
                        pass
                else:
                    row_cells[i].text = str(val)

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Legacy aliases for backward compatibility
_generate_dev_report_docx = generate_dev_report_docx
_generate_eval_report_docx = generate_eval_report_docx
_generate_ranking_report_docx = generate_ranking_report_docx
