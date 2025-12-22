# tanml/report/report_builder.py
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Inches, Mm
from pathlib import Path
import os, re, math, copy as pycopy
import datetime
from importlib.resources import files
import numpy as np  # needed for rounding helpers etc.
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL





TMP_DIR = Path(__file__).resolve().parents[1] / "tmp_report_assets"
TMP_DIR.mkdir(exist_ok=True)

class AttrDict(dict):
    def __getattr__(self, item):
        try:
            return self[item]
        except KeyError:
            raise AttributeError(item)

def _read_csv_as_rows(path, max_rows=None):
    import csv
    rows = []
    if not path or not os.path.exists(path):
        return rows
    with open(path, newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for i, row in enumerate(r):
            if max_rows is not None and i >= max_rows:
                break
            rows.append(row)
    return rows


USE_TEXT_PLACEHOLDERS = True
TXT = {
    "not_applicable": "Not applicable",
    "not_provided": "Dataset not provided",
    "none_detected": "None detected",
    "no_issues": "None (no issues detected)",
    "unknown": "Unknown",
    "dash": "—",
}
def _p(key: str) -> str:
    return TXT["dash"] if not USE_TEXT_PLACEHOLDERS else TXT.get(key, TXT["unknown"])


def _fallback_pairs_from_corr_matrix(corr_csv_path: str, threshold: float, top_k: int = 200):
    """
    Build 'top pairs' rows from a saved correlation MATRIX CSV when the engine
    did not emit 'top_pairs_main_csv'/'top_pairs_csv'.

    Returns a list of dicts with headers:
      ["feature_i", "feature_j", "corr", "n_used", "pct_missing_i", "pct_missing_j"]
    """
    rows = []
    if not corr_csv_path or not os.path.exists(corr_csv_path):
        return rows
    try:
        import pandas as pd
        df = pd.read_csv(corr_csv_path, index_col=0)
        # Ensure square matrix with aligned labels
        if df.shape[0] == 0 or df.shape[0] != df.shape[1]:
            return rows
        cols = list(df.columns)

        # Collect upper triangle |r| >= threshold
        pairs = []
        for i in range(len(cols)):
            for j in range(i + 1, len(cols)):
                try:
                    r = float(df.iloc[i, j])
                except Exception:
                    continue
                if abs(r) >= float(threshold):
                    pairs.append((cols[i], cols[j], r))

        # Sort by absolute correlation desc, then take top_k
        pairs.sort(key=lambda t: abs(t[2]), reverse=True)
        pairs = pairs[: int(top_k)]

        # Map to expected schema; we don't have n_used / pct_missing_* here
        for a, b, r in pairs:
            rows.append({
                "feature_i": a,
                "feature_j": b,
                "corr": f"{float(r):.4f}",
                "n_used": "",               # unknown in matrix-only fallback
                "pct_missing_i": "",        # unknown
                "pct_missing_j": "",        # unknown
            })
        return rows
    except Exception:
        return []

def build_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = str(h)
    for r in rows:
        vals = [r.get(h, "") for h in headers] if isinstance(r, dict) else list(r)
        vals += [""] * (len(headers) - len(vals))
        row = tbl.add_row().cells
        for i, v in enumerate(vals):
            row[i].text = str(v)
    return tbl

# --- Pretty Decile Lift Table (3-line headers, auto-fit, numeric alignment) ---
def build_decile_lift_table(doc, headers, rows):
    """
    Renders the decile lift table with:
      - compact labels
      - up to THREE lines per header cell (forced line breaks)
      - widths auto-scaled to page text width
      - right-aligned numeric columns
    """

    # 1) Compact labels (keep Total)
    label_map = {
        "decile": "Decile",
        "total": "Total",
        "events": "Events",
        "avg_score": "Avg score",
        "event_rate": "Event rate",
        "lift": "Lift",
        "cum_events": "Cum. events",
        "cum_total": "Cum. total",
        "cum_capture_rate": "Cum. capture rate",
        "cum_population": "Cum. population",
        "cum_gain": "Cum. gain",
    }
    pretty_headers = [label_map.get(h.strip(), h.replace("_", " ").strip()) for h in headers]

    # 2) Helper: force up to 3 lines per label (balanced by words)
    def split_to_max_lines(label: str, max_lines: int = 3):
        parts = str(label).split()
        if len(parts) <= 1:
            return [label]
        # Greedy balance into up to max_lines buckets
        buckets = [[] for _ in range(max_lines)]
        # pre-fill minimal distribution
        for i, w in enumerate(parts):
            buckets[i % max_lines].append(w)
        # join each bucket; trim empty
        lines = [" ".join(b).strip() for b in buckets if b]
        # remove trailing empties
        lines = [ln for ln in lines if ln]
        return lines[:max_lines]

    # 3) Build table (single header row; we’ll insert line breaks in each header cell)
    tbl = doc.add_table(rows=1, cols=len(pretty_headers))
    tbl.style = "Table Grid"
    tbl.autofit = False

    # 4) Header row with forced breaks (up to 3 lines)
    hdr_row = tbl.rows[0]
    for j, h in enumerate(pretty_headers):
        cell = hdr_row.cells[j]
        # Clear cell safely
        cell.text = ""
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        lines = split_to_max_lines(h, max_lines=3)
        if not lines:
            continue
        run = p.add_run(str(lines[0]))
        for seg in lines[1:]:
            run.add_break()            # hard line break
            p.add_run(str(seg))

    # 5) Body rows (preserve original header order)
    for r in rows:
        vals = [r.get(h, "") for h in headers] if isinstance(r, dict) else list(r)
        vals += [""] * (len(pretty_headers) - len(vals))
        cells = tbl.add_row().cells
        for j, v in enumerate(vals):
            cells[j].text = "" if v is None else str(v)

    # 6) Base widths (tight) in inches — will be scaled to fit page text width
    #    Order corresponds to incoming headers: decile, total, events, ...
    base_widths = [0.50, 0.60, 0.60, 0.70, 0.70, 0.55, 0.80, 0.80, 0.95, 0.95, 0.85]
    widths = base_widths[:len(pretty_headers)]

    # 7) Compute usable width from document section (with a small safety margin)
    try:
        sec = doc.sections[0]
        usable_in = float(sec.page_width - sec.left_margin - sec.right_margin) / 914400.0
        usable_in = max(usable_in - 0.05, 5.8)
    except Exception:
        usable_in = 6.7

    # 8) Auto-scale to fit
    total_w = sum(widths)
    if total_w > 0 and usable_in > 0:
        scale = min(1.0, usable_in / total_w)
        widths = [w * scale for w in widths]

    # 9) Apply widths
    for j, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[j].width = Inches(w)

    # 10) Right-align numeric columns
    numeric_like = {
        "Total", "Events", "Avg score", "Event rate", "Lift",
        "Cum. events", "Cum. total", "Cum. capture rate",
        "Cum. population", "Cum. gain",
    }
    for i, h in enumerate(pretty_headers):
        right = h in numeric_like
        for row in tbl.rows[1:]:
            for p in row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.LEFT

    return tbl



def insert_after(doc, anchor, tbl):
    for p in doc.paragraphs:
        if anchor.lower() in p.text.lower():
            parent = p._p.getparent()
            parent.insert(parent.index(p._p) + 1, tbl._tbl)
            return
    print(f"⚠️ anchor «{anchor}» not found")

def insert_image_after(doc, anchor, img_path, width_in=6.5):
    """
    Inserts a single image after the paragraph containing 'anchor'.
    """
    from docx.shared import Inches
    for p in doc.paragraphs:
        if anchor.lower() in p.text.lower():
            # Create new paragraph
            new_p = doc.add_paragraph() 
            run = new_p.add_run()
            try:
                run.add_picture(img_path, width=Inches(width_in))
                # Move it after 'p'
                p._p.addnext(new_p._p)
            except Exception as e:
                print(f"Failed to add picture {img_path}: {e}")
            return
    print(f"⚠️ anchor «{anchor}» not found (for image)")

def insert_image_grid(doc, anchor: str, img_paths, cols: int = 3, width_in: float = 2.2):
    paths = [p for p in (img_paths or []) if p and os.path.exists(p)]
    if not paths:
        for p in doc.paragraphs:
            if anchor.lower() in p.text.lower():
                # Correctly insert paragraph AFTER 'p'
                new_p = p._parent.add_paragraph()
                new_p.add_run("(no plots available)")
                p._p.addnext(new_p._p)
                return
        print(f"⚠️ anchor «{anchor}» not found (no plots)")
        return

    rows = max(1, math.ceil(len(paths) / max(1, cols)))
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    i = 0
    for r in range(rows):
        for c in range(cols):
            cell = tbl.cell(r, c)
            if i < len(paths):
                para = cell.paragraphs[0]
                para.text = ""
                run = para.add_run()
                run.add_picture(paths[i], width=Inches(width_in))
                i += 1
            else:
                cell.text = ""
                cell.text = ""
    insert_after(doc, anchor, tbl)

def _log_debug(msg):
    with open("/tmp/tanml_report_debug.log", "a") as f:
        f.write(f"{datetime.datetime.now()} - {msg}\n")

def _replace_placeholder_with_kv_table(doc, placeholder: str, data: dict):
    """
    Finds 'placeholder' in doc (paragraphs or table cells) and replaces it with
    a 2-column key-value table.
    """
    if not data or not isinstance(data, dict):
        return False

    # Helper to populate table
    def _fill(tbl):
        # tbl.style = "Table Grid"  # Optional: might inherit from parent if nested
        for k, v in data.items():
            row = tbl.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

    # 2. Search in existing tables (including nested tables)
    def _traverse_tables(tables_list, depth=0):
        for t_idx, t in enumerate(tables_list):
            for r_idx, row in enumerate(t.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Check strictly inside this cell's direct text first
                    # We try to find it in paragraphs
                    for p in cell.paragraphs:
                        # Use regex to be robust against spaces: {{ ModelMetaCheck.target_balance }}
                        # Regex pattern: {{ ?ModelMetaCheck\.target_balance ?}}
                        if re.search(r"\{\{\s*ModelMetaCheck\.target_balance\s*\}\}", p.text):
                            # Remove placeholder using regex sub
                            p.text = re.sub(r"\{\{\s*ModelMetaCheck\.target_balance\s*\}\}", "", p.text)
                            
                            # Create nested table here
                            try:
                                nested = cell.add_table(rows=0, cols=2)
                                nested.style = "Table Grid"
                                _fill(nested)
                                # Move nested table after the modified paragraph
                                p._p.addnext(nested._tbl)
                            except Exception as e:
                                print(f"Failed to add nested table: {e}")
                                # Fallback
                                for k, v in data.items():
                                    cell.add_paragraph(f"{k}: {v}")
                            return True
                    
                    # Recurse into nested tables
                    if cell.tables:
                         if _traverse_tables(cell.tables, depth=depth+1):
                             return True
        return False

    # 1. Search in body paragraphs
    for p_idx, p in enumerate(doc.paragraphs):
        # Use regex to be robust against spaces: {{ ModelMetaCheck.target_balance }}
        # Regex pattern: {{ ?ModelMetaCheck\.target_balance ?}}
        if re.search(r"\{\{\s*ModelMetaCheck\.target_balance\s*\}\}", p.text):
            p.text = re.sub(r"\{\{\s*ModelMetaCheck\.target_balance\s*\}\}", "", p.text)
            
            # Create isolated table
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Table Grid"
            _fill(tbl)
            # Move after p
            p._p.addnext(tbl._tbl)
            return True
            
    return _traverse_tables(doc.tables)

# ---------- placeholder & image replacement ----------
_PLACEHOLDER = re.compile(r"\{\{\s*([A-Za-z0-9_\.]+)\s*\}\}")
_IMG_MARKER  = re.compile(r"\[\[IMG:([A-Za-z0-9_\-]+)\]\]")

def _get_nested(d, dotted, default=""):
    cur = d
    for part in dotted.split("."):
        if not isinstance(cur, dict) or part not in cur:
            return default
        cur = cur[part]
    return cur

def _replace_text_placeholders(doc, mapping):
    # paragraphs
    for p in doc.paragraphs:
        full = p.text
        def repl(m):
            key = m.group(1)
            return str(mapping.get(key, _get_nested(mapping, key, "")) or "")
        new_full = _PLACEHOLDER.sub(repl, full)
        if new_full != full:
            for _ in range(len(p.runs)-1, -1, -1):
                p.runs[_].text = ""
            if p.runs:
                p.runs[0].text = new_full
            else:
                p.add_run(new_full)
    # table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = p.text
                    new_full = _PLACEHOLDER.sub(
                        lambda m: str(mapping.get(m.group(1), _get_nested(mapping, m.group(1), "")) or ""),
                        full
                    )
                    if new_full != full:
                        for _ in range(len(p.runs)-1, -1, -1):
                            p.runs[_].text = ""
                        if p.runs:
                            p.runs[0].text = new_full
                        else:
                            p.add_run(new_full)

def _replace_image_markers(doc, images_map, *, width_in=4.8):
    def _put_image_in_paragraph(p, path):
        if path and os.path.exists(path):
            p.text = ""
            run = p.add_run()
            run.add_picture(path, width=Inches(width_in))
        else:
            p.text = "(image not available)"
    # paragraphs
    for p in list(doc.paragraphs):
        m = _IMG_MARKER.search(p.text)
        if m:
            key = m.group(1)
            _put_image_in_paragraph(p, images_map.get(key))
    # table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    m = _IMG_MARKER.search(p.text)
                    if m:
                        key = m.group(1)
                        _put_image_in_paragraph(p, images_map.get(key))

# ---------- formatters ----------
def _fmt2(x, nd=2, dash="—"):
    try:
        if x is None: return dash
        xf = float(x)
        if xf != xf:  # NaN
            return dash
        return f"{xf:.{nd}f}"
    except Exception:
        return dash

def _fmt_ratio_or_pct(x, pct=False, dash="—", nd=2):
    try:
        if x is None:
            return dash
        xf = float(x)
        if pct:
            return f"{xf*100:.{nd}f}%"
        return f"{xf:.{nd}f}"
    except Exception:
        return dash

def _fmt_list(lst, *, max_items=20, sep=", "):
    if not lst:
        return "—"
    if isinstance(lst, (str, bytes)):
        return str(lst)
    try:
        seq = list(lst)
    except Exception:
        return str(lst)
    if len(seq) <= max_items:
        return sep.join(map(str, seq))
    head = sep.join(map(str, seq[:max_items]))
    return f"{head}{sep}… (+{len(seq)-max_items} more)"


# put this right below _fmt_list(...)
def _fmt_list_or_message(lst, *, empty_msg="None (no issues detected)", max_items=20, sep=", "):
    if not lst:
        return empty_msg
    try:
        seq = list(lst) if not isinstance(lst, (str, bytes)) else [lst]
    except Exception:
        return str(lst)
    if len(seq) <= max_items:
        return sep.join(map(str, seq))
    head = sep.join(map(str, seq[:max_items]))
    return f"{head}{sep}… (+{len(seq)-max_items} more)"


def _fmt_feature_names(v, *, max_names=30):
    if v is None:
        return ""
    if isinstance(v, (list, tuple)):
        if len(v) <= max_names:
            return ", ".join(map(str, v))
        head = ", ".join(map(str, v[:max_names]))
        return f"{head}, … (+{len(v)-max_names} more)"
    return str(v)

def _fmt_target_balance(tb):
    if not isinstance(tb, dict) or not tb:
        return ""
    vals = list(tb.values())
    are_probs = all(isinstance(x, (int, float)) and 0 <= float(x) <= 1 for x in vals)
    items = []
    for k in sorted(tb.keys()):
        v = tb[k]
        if are_probs:
            try:
                items.append(f"{k}: {float(v)*100:.1f}%")
            except Exception:
                items.append(f"{k}: {v}")
        else:
            items.append(f"{k}: {v}")
    return ", ".join(items)

# ---------- SHAP helpers ----------
def _attach_shap_to_context(results: dict, context: dict) -> None:
    shap_ctx = {
        "shap_section": False,
        "shap_beeswarm_path": None,
        "shap_bar_path": None,
        "shap_top_features": [],
    }
    try:
        shap_res = (results or {}).get("SHAPCheck") or {}
        if "SHAPCheck" in shap_res and isinstance(shap_res["SHAPCheck"], dict):
            shap_res = shap_res["SHAPCheck"]
        plots = shap_res.get("plots") or {}
        bees = plots.get("beeswarm")
        barp = plots.get("bar")
        legacy = shap_res.get("shap_plot_path")
        if not bees and legacy:
            bees = legacy
        shap_ctx["shap_beeswarm_path"] = bees if bees and os.path.exists(bees) else None
        shap_ctx["shap_bar_path"] = barp if barp and os.path.exists(barp) else None
        shap_ctx["shap_top_features"] = shap_res.get("top_features") or []
        shap_ctx["shap_section"] = bool(
            shap_ctx["shap_beeswarm_path"] or shap_ctx["shap_bar_path"] or shap_ctx["shap_top_features"]
        )
    except Exception as e:
        print("⚠️ SHAP context attach failed:", e)
    context.update(shap_ctx)

# --- main class --------------------------------------------------------------
class ReportBuilder:
    """Build a Word report from validation results (no docxtpl)."""

    def __init__(self, results, template_path, output_path):
        self.results = results or {}
        self.template_path = template_path or files("tanml.report.templates").joinpath("report_template.docx")
        self.output_path = output_path

        corr = self.results.get("CorrelationCheck", {}) or {}
        corr_artifacts = corr.get("artifacts", corr)
        self.corr_heatmap_path = corr_artifacts.get("heatmap_path")
        self.corr_pearson_path = corr_artifacts.get("pearson_csv", "N/A")
        self.corr_spearman_path = corr_artifacts.get("spearman_csv", "N/A")

    def _grab(self, name, default=None):
        return (
            self.results.get(name)
            or self.results.get("check_results", {}).get(name)
            or default
        )

    @staticmethod
    def _extract_baseline_logit_rows(ctx):
        """
        Read Logit baseline metrics from ctx["LogitStats"]["baseline_metrics"]
        and map to a simple 2-col table.
        """
        logit = (ctx.get("LogitStats") or {})
        bm = logit.get("baseline_metrics") or {}
        summary = bm.get("summary") or bm  # support both shapes

        order = [
            ("ROC-AUC",      ("AUC","auc","roc_auc","rocauc")),
            ("KS",           ("KS","ks")),
            ("F1",           ("F1","f1","f1_score")),
            ("PR AUC",       ("PR_AUC","pr_auc","average_precision","prauc")),
            ("Gini",         ("GINI","gini")),
            ("Precision",    ("Precision","precision","prec")),
            ("Recall",       ("Recall","recall","tpr","sensitivity")),
            ("Accuracy",     ("Accuracy","accuracy","acc")),
            ("Brier Score",  ("Brier","brier","brier_score")),
            ("Log Loss",     ("Log Loss","log_loss","logloss")),
        ]
        rows = []
        for label, keys in order:
            for k in keys:
                if k in summary and summary[k] is not None:
                    v = summary[k]
                    try:
                        v = f"{float(v):.4f}"
                    except Exception:
                        v = str(v)
                    rows.append({"metric": label, "value": v})
                    break
        return rows

    def build(self):
        # ===== 1) Build context =================================================
        self.results.setdefault("RuleEngineCheck", AttrDict({"rules": {}, "overall_pass": True}))
        ctx = pycopy.deepcopy(self.results)
        ctx.update(self.results.get("check_results", {}))

        # unwrap {"Key":{"Key":...}}
        for k, v in list(ctx.items()):
            if isinstance(v, dict) and k in v and len(v) == 1:
                ctx[k] = v[k]

        # defaults
        for k, note in [
            ("RawDataCheck", "Raw-data check skipped"),
            ("CleaningReproCheck", "Cleaning-repro check skipped"),
        ]:
            ctx.setdefault(k, AttrDict({"note": note}))

        # Model meta fallback
        if "ModelMetaCheck" not in ctx or "model_type" not in ctx["ModelMetaCheck"]:
            meta_fields = [
                "model_type", "model_class", "module", "n_features", "feature_names",
                "n_train_rows", "target_balance", "hyperparam_table", "attributes",
            ]
            meta = {f: self.results.get(f) for f in meta_fields if self.results.get(f) is not None}
            ctx["ModelMetaCheck"] = AttrDict(meta or {"note": "Model metadata not available"})

        # SHAP
        _attach_shap_to_context(self.results, ctx)

        # EDA
        eda = self._grab("EDACheck", {})
        ctx["eda_summary_path"] = eda.get("summary_stats", "N/A")
        ctx["eda_missing_path"] = eda.get("missing_values", "N/A")
        # Helper to resolve path
        def _resolve_eda_path(fn):
            # If it's already a valid path, use it
            if os.path.exists(fn):
                return fn
            # Otherwise try joining with reports/eda
            joined = os.path.join("reports/eda", fn)
            if os.path.exists(joined):
                return joined
            return None

        ctx["eda_images_paths"] = [
            p for p in [
                _resolve_eda_path(fn) 
                for fn in (eda.get("visualizations", []) or [])
            ] 
            if p is not None
        ]

        # Correlation
        corr_res = self._grab("CorrelationCheck", {}) or {}
        corr_artifacts = corr_res.get("artifacts", corr_res)
        corr_summary = corr_res.get("summary", {}) or {}

        def _pick(d, *keys, default=None):
            for k in keys:
                if k in d and d[k] is not None:
                    return d[k]
            return default

        top_pairs_csv = (
            corr_artifacts.get("top_pairs_csv")
            or corr_artifacts.get("pairs_csv")
            or corr_artifacts.get("csv")
        )
        top_pairs_main_csv = (
            corr_artifacts.get("top_pairs_main_csv")
            or corr_artifacts.get("main_csv")
        )
        heatmap_path = (
            corr_artifacts.get("heatmap_path")
            or corr_artifacts.get("heatmap")
            or self.corr_heatmap_path
        )

        method = _pick(corr_summary, "method", "corr_method")
        threshold = _pick(corr_summary, "threshold", "high_corr_threshold")
        n_numeric_features = _pick(corr_summary, "n_numeric_features", "numeric_features", "n_features_numeric")
        plotted_features = _pick(corr_summary, "plotted_features", "features_plotted")
        plotted_full_matrix = _pick(corr_summary, "plotted_full_matrix", "full_matrix")
        n_pairs_total = _pick(corr_summary, "n_pairs_total", "pairs_total")
        n_pairs_flagged = _pick(corr_summary, "n_pairs_flagged_ge_threshold", "n_pairs_flagged")

        ctx.setdefault("CorrelationCheck", {})
        ctx["CorrelationCheck"].update({
            "method": method,
            "threshold": threshold,
            "n_numeric_features": n_numeric_features,
            "plotted_features": plotted_features,
            "plotted_full_matrix": plotted_full_matrix,
            "n_pairs_total": n_pairs_total,
            "n_pairs_flagged_ge_threshold": n_pairs_flagged,
            "top_pairs_csv": top_pairs_csv,
            "top_pairs_main_csv": top_pairs_main_csv,
        })
        ctx["correlation_pearson_path"] = self.corr_pearson_path
        ctx["correlation_spearman_path"] = self.corr_spearman_path
        ctx["correlation_heatmap_path"] = heatmap_path

        corr_preview_headers = [
            "feature_i", "feature_j", "corr",
            "n_used", "pct_missing_i", "pct_missing_j",
        ]
        pairs_csv_for_preview = top_pairs_main_csv or top_pairs_csv
        corr_preview_rows = _read_csv_as_rows(pairs_csv_for_preview, max_rows=None)

        # round numeric fields — esp. 'corr' to 4 decimals
        rounded_corr_rows = []
        for r in corr_preview_rows:
            new_row = {}
            for h in corr_preview_headers:
                val = r.get(h, "")
                if h == "corr":
                    try:
                        new_row[h] = f"{float(val):.4f}"
                    except Exception:
                        new_row[h] = val
                elif h in ("pct_missing_i", "pct_missing_j"):
                    try:
                        new_row[h] = f"{float(val):.2f}"
                    except Exception:
                        new_row[h] = val
                else:
                    new_row[h] = val
            rounded_corr_rows.append(new_row)

        corr_preview_rows = rounded_corr_rows



        # ---- Fallback: if no top-pairs CSV rows, derive from pearson/spearman matrix ----
        if not corr_preview_rows:
            cc_summary = (ctx.get("CorrelationCheck") or {}).get("summary", {}) or {}
            thr = cc_summary.get("threshold")
            if thr is None:
                # Also try the normalized keys we stored earlier
                thr = (ctx.get("CorrelationCheck") or {}).get("threshold") or 0.80

            # Prefer pearson; fall back to spearman
            pearson_path = ctx.get("correlation_pearson_path")
            spearman_path = ctx.get("correlation_spearman_path")

            fallback_rows = []
            if pearson_path and os.path.exists(pearson_path):
                fallback_rows = _fallback_pairs_from_corr_matrix(pearson_path, float(thr), top_k=200)
            elif spearman_path and os.path.exists(spearman_path):
                fallback_rows = _fallback_pairs_from_corr_matrix(spearman_path, float(thr), top_k=200)

            # Use fallback if we found any
            if fallback_rows:
                corr_preview_rows = fallback_rows


        # ---------- Performance (classification) ----------
        perf_root = self.results.get("performance", {}) or {}
        cls = perf_root.get("classification", {}) or {}
        cls_summary = (cls.get("summary") or {})
        cls_plots = (cls.get("plots") or {})
        cls_tables = (cls.get("tables") or {})

        def _pick_first(d, *keys):
            for k in keys:
                if k in d and d[k] is not None:
                    return d[k]
            low = {str(k).lower(): v for k, v in d.items()}
            for k in keys:
                lk = str(k).lower()
                if lk in low and low[lk] is not None:
                    return low[lk]
            return None

        ctx["classification_summary"] = {
            "AUC":       _pick_first(cls_summary, "AUC", "auc", "roc_auc"),
            "KS":        _pick_first(cls_summary, "KS", "ks"),
            "GINI":      _pick_first(cls_summary, "GINI", "gini"),
            "PR_AUC":    _pick_first(cls_summary, "PR_AUC", "pr_auc", "average_precision"),
            "F1":        _pick_first(cls_summary, "F1", "f1", "f1_score"),
            "Precision": _pick_first(cls_summary, "Precision", "precision"),
            "Recall":    _pick_first(cls_summary, "Recall", "recall", "tpr", "sensitivity"),
            "Accuracy":  _pick_first(cls_summary, "Accuracy", "accuracy"),
            "Brier":     _pick_first(cls_summary, "Brier", "brier", "brier_score"),
        }
        ctx["classification_tables"] = {
            "confusion": _read_csv_as_rows(_pick_first(cls_tables, "confusion_csv", "confusion")),
            "lift":      _read_csv_as_rows(_pick_first(cls_tables, "lift_csv", "decile_lift_csv", "gain_lift_csv")),
        }
        ctx["classification_plot_paths"] = {
            "roc":         _pick_first(cls_plots, "roc", "roc_curve"),
            "pr":          _pick_first(cls_plots, "pr", "pr_curve", "precision_recall"),
            "lift":        _pick_first(cls_plots, "lift", "cumulative_gain", "lift_curve"),
            "calibration": _pick_first(cls_plots, "calibration", "reliability"),
            "confusion":   _pick_first(cls_plots, "confusion", "confusion_matrix"),
            "ks":          _pick_first(cls_plots, "ks", "ks_curve"),
        }

        # Regression metrics
        reg = self._grab("RegressionMetrics", {}) or {}
        reg_art = reg.get("artifacts", {}) or {}
        ctx.setdefault("RegressionMetrics", {})
        ctx["RegressionMetrics"].update({
            "notes": reg.get("notes", []),
            "rmse": reg.get("rmse"),
            "mae": reg.get("mae"),
            "median_ae": reg.get("median_ae"),
            "r2": reg.get("r2"),
            "r2_adjusted": reg.get("r2_adjusted"),
            "mape_or_smape": reg.get("mape_or_smape"),
            "mape_used": reg.get("mape_used"),
            "artifacts": {
                "pred_vs_actual": reg_art.get("pred_vs_actual"),
                "residuals_vs_pred": reg_art.get("residuals_vs_pred"),
                "residual_hist": reg_art.get("residual_hist"),
                "qq_plot": reg_art.get("qq_plot"),
                "abs_error_box": reg_art.get("abs_error_box"),
                "abs_error_violin": reg_art.get("abs_error_violin"),
            }
        })

        # Summary (ensure)
        ctx.setdefault("summary", {})
        ctx["summary"].setdefault("rmse", ctx["RegressionMetrics"].get("rmse"))
        ctx["summary"].setdefault("mae",  ctx["RegressionMetrics"].get("mae"))
        ctx["summary"].setdefault("r2",   ctx["RegressionMetrics"].get("r2"))
        if "task_type" not in ctx:
            ctx["task_type"] = "regression" if ctx["RegressionMetrics"].get("rmse") is not None else "classification"

        # Stress list→dict back-compat
        if isinstance(self.results.get("StressTestCheck"), list):
            ctx["StressTestCheck"] = {"table": self.results["StressTestCheck"]}

        # Input Cluster Coverage
        icc = self._grab("InputClusterCoverageCheck", {}) or self._grab("InputClusterCheck", {}) or {}
        icc_art = icc.get("artifacts", icc)
        cluster_rows = icc.get("cluster_table") or icc_art.get("cluster_table") or []
        cluster_csv = (icc.get("cluster_csv") or icc_art.get("cluster_csv") or icc_art.get("csv") or "—")
        plot_path = (icc.get("cluster_plot_img") or icc_art.get("cluster_plot_img") or icc_art.get("plot_img"))

        ctx.setdefault("InputClusterCheck", {})
        ctx["InputClusterCheck"]["cluster_csv"] = cluster_csv
        ctx["InputClusterCheck"]["cluster_table"] = [
            {
                "Cluster": r.get("Cluster") or r.get("cluster") or r.get("label"),
                "Count": r.get("Count") or r.get("count"),
                "Percent": r.get("Percent") or r.get("percent"),
            }
            for r in cluster_rows if isinstance(r, dict)
        ]
        ctx["ks_curve_path"] = ctx["classification_plot_paths"].get("ks")

        # VIF normalize
        vif = self._grab("VIFCheck", {})
        ctx["VIFCheck"] = AttrDict(
            vif if isinstance(vif, dict) and "vif_table" in vif
            else {"vif_table": [], "high_vif_features": [], "error": "Invalid VIFCheck"}
        )

        # ===== 2) Open template, replace text placeholders & images ==========
        # Use DocxTemplate for Jinja2 conditionals ({% if %}, etc.)
        # Add use_cv flag to context for conditional rendering
        cv_stats = ctx.get("summary", {}).get("cv_stats", {})
        ctx["use_cv"] = bool(cv_stats)  # True if CV was run
        
        # For docxtpl: we only need to process {% if %} conditionals
        # Other {{placeholders}} will be handled by _replace_text_placeholders later
        # Custom undefined class that renders undefined variables back to {{var}} syntax
        from jinja2 import Environment, Undefined
        
        class PassthroughUndefined(Undefined):
            def __str__(self):
                return '{{' + self._undefined_name + '}}'
            
            def __getattr__(self, name):
                # For nested attributes like ModelMetaCheck.target_balance
                return PassthroughUndefined(name=f'{self._undefined_name}.{name}')
        
        jinja_env = Environment(undefined=PassthroughUndefined)
        tpl = DocxTemplate(str(self.template_path))
        
        # --- Prepare Jinja2 Context for docxtpl ---
        # We must provide keys for ALL {% if %} blocks and {{ variables }} handled by docxtpl.
        jinja_ctx = {"use_cv": ctx["use_cv"]}

        # 1. EDA Placeholders (must pass to render so they aren't stripped/broken)
        jinja_ctx["eda_summary_path"] = ctx.get("eda_summary_path") or "(not available)"
        jinja_ctx["eda_missing_path"] = ctx.get("eda_missing_path") or "(not available)"
        
        # Replicate EDA count note logic
        all_eda = ctx.get("eda_images_paths") or []
        opts = ctx.get("report_options") or {}
        max_p = opts.get("max_plots", -1)
        if max_p in (-1, None):
            eda_subset = all_eda
        else:
            eda_subset = all_eda[:int(max_p)]
        jinja_ctx["eda_count_note"] = f"(showing {len(eda_subset)} of {len(all_eda)})" if len(all_eda) != len(eda_subset) else ""

        # 2. CV Regression Flags
        # Optimistically set True if we are in CV Regression mode, so {% if %} blocks render.
        # The actual images will be filled in by _replace_image_markers later.
        oof = cv_stats.get("oof", {})
        if ctx["use_cv"] and "y_pred" in oof:
            for k in [
                "cv_reg_residuals_vs_pred", "cv_reg_residual_hist", 
                "cv_reg_qq", "cv_reg_abs_error_box", "cv_reg_abs_error_violin"
            ]:
                jinja_ctx[k] = True

        try:

            tpl.render(jinja_ctx, jinja_env=jinja_env)  # Render logic

            doc = tpl  # DocxTemplate is a Document subclass
        except Exception as e:
            print(f"Warning: docxtpl render failed: {e}, falling back to Document")
            doc = Document(str(self.template_path))

        # -------- scalar_map (text placeholders) ----------------------------
        from datetime import datetime, timezone
        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        
        # Determine model name
        mm = ctx.get("ModelMetaCheck", {}) or {}
        m_name = mm.get("model_class") or "Model"
        
        scalar_map = {
            "generated_on": now_str,
            "validation_date": now_str,
            "validated_by": "TanML",
            "model_file": f"{m_name}.pkl",
            "model_path": ctx.get("model_path", f"{m_name}.pkl"),
            "task_type": (ctx.get("task_type") or "classification").title(),
            "ModelMetaCheck.model_class":  mm.get("model_class", ""),   
            "ModelMetaCheck.module":       (ctx.get("ModelMetaCheck", {}) or {}).get("module", ""),        
            "ModelMetaCheck.model_type":   (ctx.get("ModelMetaCheck", {}) or {}).get("model_type", ""),
            "ModelMetaCheck.n_features":   (ctx.get("ModelMetaCheck", {}) or {}).get("n_features", ""),
            "ModelMetaCheck.feature_names": _fmt_feature_names((ctx.get("ModelMetaCheck", {}) or {}).get("feature_names")),
            "ModelMetaCheck.n_train_rows": (ctx.get("ModelMetaCheck", {}) or {}).get("n_train_rows", ""),
        }

        # --- Logistic (Logit) summary text for classification template ---
        logit_ctx = ctx.get("LogitStats") or {}
        scalar_map["LogitStats.summary_text"] = logit_ctx.get("summary_text") or ""

        # REGRESSION: Target Balance Table Replacement
        # If we have the regression summary dict (Range/Mean/Std), inject a table instead of text
        mb_data = (ctx.get("ModelMetaCheck", {}) or {}).get("target_balance")
        
        _log_debug(f"Target Balance Check: mb_data type={type(mb_data)}")
        if isinstance(mb_data, dict):
             _log_debug(f"mb_data keys: {list(mb_data.keys())}")
        
        # Relaxed check: just look for 'Mean' to be safe, or even just dict
        if isinstance(mb_data, dict) and "Mean" in mb_data:
            _log_debug("Entering table replacement logic...")
            try:
                # Attempt table replacement
                replaced = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", mb_data)
                _log_debug(f"Replacement result: {replaced}")
                if replaced:
                    # Prevent text replacement from re-processing it (though it's gone)
                    scalar_map["ModelMetaCheck.target_balance"] = ""
                else:
                    _log_debug("Replacement returned False. Fallback to text.")
                    items = [f"{k}: {v}" for k, v in mb_data.items()]
                    scalar_map["ModelMetaCheck.target_balance"] = "\n".join(items)
            except Exception as e:
                _log_debug(f"Error during replacement: {e}")
        else:
            _log_debug("Skipping table replacement: mb_data missing Mean/Std or not dict")
        
        # REGRESSION rounded values / labels
        scalar_map.update({
            "summary.rmse2": _fmt2(ctx.get("summary", {}).get("rmse")),
            "summary.mae2":  _fmt2(ctx.get("summary", {}).get("mae")),
            "summary.r22":   _fmt2(ctx.get("summary", {}).get("r2")),
            "RegressionMetrics.r2_adjusted2": _fmt2(ctx.get("RegressionMetrics", {}).get("r2_adjusted")),
            "RegressionMetrics.median_ae2":   _fmt2(ctx.get("RegressionMetrics", {}).get("median_ae")),
        })

        scalar_map.update({
            "eda_summary_path": ctx.get("eda_summary_path") or "(not available)",
            "eda_missing_path": ctx.get("eda_missing_path") or "(not available)",
        })
        rm = ctx.get("RegressionMetrics", {}) or {}
        mape_or_smape = rm.get("mape_or_smape")
        mape_used = rm.get("mape_used")
        scalar_map["RegressionMetrics.mape_label"] = (
            "N/A" if mape_or_smape is None else f"{_fmt2(mape_or_smape)}% ({'MAPE' if mape_used else 'SMAPE'})"
        )
        notes_list = rm.get("notes") or []
        notes_text = "\n".join(map(str, notes_list)).strip() if notes_list else ""
        scalar_map["RegressionMetrics.notes_text"] = notes_text if notes_text else "None"

        # CLASSIFICATION rounded values (Simple vs CV Robust)
        cs = ctx.get("classification_summary", {}) or {}
        cv_stats = ctx.get("summary", {}).get("cv_stats", {})
        
        def _get_metric(key, simple_key):
            # Try to grab CV Mean ± Std first
            if key in cv_stats:
                m = cv_stats[key]
                return f"{m['mean']:.2f} ± {m['std']:.2f}"
            # Fallback to simple
            return _fmt2(cs.get(simple_key))

        scalar_map.update({
            "classification_summary.AUC2":       _get_metric("roc_auc", "AUC"),
            "classification_summary.KS2":        _fmt2(cs.get("KS")), # CV for KS not implemented in app loop yet
            "classification_summary.F12":        _get_metric("f1", "F1"),
            "classification_summary.PR_AUC2":    _fmt2(cs.get("PR_AUC")),
            "classification_summary.GINI2":      _fmt2(cs.get("GINI")),
            "classification_summary.Precision2": _get_metric("precision", "Precision"),
            "classification_summary.Recall2":    _get_metric("recall", "Recall"),
            "classification_summary.Accuracy2":  _get_metric("accuracy", "Accuracy"),
            "classification_summary.Brier2":     _fmt2(cs.get("Brier")),
        })

        # --- CV Detailed Stats (Repeated CV) ---
        # Maps cv_stats["metric"]["stat"] -> {{cv.metric.stat}}
        # e.g. {{cv.auc.mean}}, {{cv.rmse.p95}}
        cv_stats_root = ctx.get("summary", {}).get("cv_stats", {})
        
        # known metrics to look for
        cv_metric_keys = [
            # Classification
            ("auc", "auc"), ("roc_auc", "auc"), 
            ("pr_auc", "pr_auc"), ("average_precision", "pr_auc"),
            ("ks", "ks"),
            ("log_loss", "logloss"), ("logloss", "logloss"),
            ("brier", "brier"), ("brier_score", "brier"),
            ("gini", "gini"),
            ("f1", "f1"), ("f1_score", "f1"),
            ("precision", "precision"),
            ("recall", "recall"), ("sensitivity", "recall"),
            ("accuracy", "accuracy"),
            ("mcc", "mcc"), ("matthews_corrcoef", "mcc"),
            ("balanced_accuracy", "bal_acc"), ("bal_acc", "bal_acc"),
            
            # Regression
            ("rmse", "rmse"), 
            ("mae", "mae"), 
            ("median_ae", "median_ae"),
            ("r2", "r2"),
            ("smape", "smape"), ("mape", "smape"), # unify
        ]
        
        # standard stats to extract
        stats_to_map = ["mean", "std", "p05", "p50", "p95", "min", "max"]

        for src_key, out_key in cv_metric_keys:
            if src_key in cv_stats_root:
                m_obj = cv_stats_root[src_key]
                if isinstance(m_obj, dict):
                    for stat in stats_to_map:
                        val = m_obj.get(stat)
                        # format logic: rounds to 4 decimals usually good for table
                        placeholder = f"cv.{out_key}.{stat}"
                        scalar_map[placeholder] = _fmt2(val, nd=4)

        # Threshold info
        thr_info = cv_stats_root.get("threshold_info") or {}
        scalar_map["cv.threshold.rule"] = str(thr_info.get("rule") or "—")
        scalar_map["cv.threshold.value"] = _fmt2(thr_info.get("value"), nd=4)

        # DataQualityCheck (train/test)
        dq = ctx.get("DataQualityCheck") or {}
        def _pick_dq(dq_root, split, key, fallback_key=None):
            split_d = dq_root.get(split) or {}
            if key in split_d and split_d[key] is not None:
                return split_d[key]
            if fallback_key and (fallback_key in dq_root) and dq_root[fallback_key] is not None:
                return dq_root[fallback_key]
            alt = {
                ("avg_missing",): ["avg_missing_rate", "missing_rate"],
                ("columns_with_missing",): ["cols_with_missing", "missing_columns"],
            }
            for k_tuple, alts in alt.items():
                if key in k_tuple:
                    for a in alts:
                        if split_d.get(a) is not None:
                            return split_d[a]
                        if fallback_key and dq_root.get(a) is not None:
                            return dq_root[a]
            return None

        train_avg_missing = _pick_dq(dq, "train", "avg_missing", fallback_key="avg_missing")
        test_avg_missing  = _pick_dq(dq, "test",  "avg_missing", fallback_key="avg_missing")
        train_cols_mis = _pick_dq(dq, "train", "columns_with_missing", fallback_key="columns_with_missing")
        test_cols_mis  = _pick_dq(dq, "test",  "columns_with_missing", fallback_key="columns_with_missing")
        const_cols_dq = dq.get("constant_columns")
        if const_cols_dq is None:
            const_cols_dq = (dq.get("train") or {}).get("constant_columns") or (dq.get("test") or {}).get("constant_columns")

        scalar_map.update({
            "DataQualityCheck.train_avg_missing": _fmt_ratio_or_pct(train_avg_missing, pct=True, dash=_p("unknown")),
            "DataQualityCheck.test_avg_missing":  _fmt_ratio_or_pct(test_avg_missing,  pct=True, dash=_p("unknown")),
            "DataQualityCheck.train_cols_missing": _fmt_list_or_message(train_cols_mis, empty_msg=_p("none_detected"), max_items=25),
            "DataQualityCheck.test_cols_missing":  _fmt_list_or_message(test_cols_mis,  empty_msg=_p("none_detected"), max_items=25),
            "DataQualityCheck.constant_columns_str": _fmt_list_or_message(const_cols_dq, empty_msg=_p("none_detected"), max_items=25),
        })

        # RawDataCheck
        rd = ctx.get("RawDataCheck") or {}
        total_rows = rd.get("total_rows") or rd.get("n_rows") or rd.get("rows")
        total_cols = rd.get("total_columns") or rd.get("n_columns") or rd.get("columns")
        avg_missing = rd.get("avg_missing") or rd.get("avg_missing_rate") or rd.get("missing_rate")
        dup_rows = rd.get("duplicate_rows") or rd.get("n_duplicate_rows") or rd.get("duplicates")
        cols_with_missing = rd.get("columns_with_missing") or rd.get("cols_with_missing") or rd.get("missing_columns")
        const_cols = rd.get("constant_columns") or rd.get("constant")
        raw_skipped = bool(rd.get("note")) and all(
            v is None for v in [total_rows, total_cols, avg_missing, dup_rows, cols_with_missing, const_cols]
        )

        scalar_map.update({
            "RawDataCheck.total_rows":
                _p("not_provided") if raw_skipped else (total_rows if total_rows is not None else _p("unknown")),
            "RawDataCheck.total_columns":
                _p("not_provided") if raw_skipped else (total_cols if total_cols is not None else _p("unknown")),
            "RawDataCheck.avg_missing_pct":
                _p("not_provided") if raw_skipped else _fmt_ratio_or_pct(avg_missing, pct=True, dash=_p("unknown")),
            "RawDataCheck.columns_with_missing_str":
                _p("not_provided") if raw_skipped else _fmt_list_or_message(cols_with_missing, empty_msg=_p("none_detected"), max_items=30),
            "RawDataCheck.duplicate_rows":
                _p("not_provided") if raw_skipped else (str(dup_rows) if dup_rows not in (None, 0) else _p("none_detected")),
            "RawDataCheck.constant_columns_str":
                _p("not_provided") if raw_skipped else _fmt_list_or_message(const_cols, empty_msg=_p("none_detected"), max_items=30),
        })


        # Correlation formatted fields & notes
        cc = (ctx.get("CorrelationCheck") or {})
        def _fmt_int(x, dash="—"):
            try:
                if x is None: return dash
                return str(int(round(float(x))))
            except Exception:
                return dash
        def _fmt_float(x, nd=2, dash="—"):
            try:
                if x is None: return dash
                return f"{float(x):.{nd}f}"
            except Exception:
                return dash

        notes_val = cc.get("notes")
        if isinstance(notes_val, (list, tuple)):
            notes_text = "Notes:\n" + "\n".join(f"- {str(n)}" for n in notes_val) if notes_val else ""
        elif isinstance(notes_val, str):
            notes_text = "Notes:\n- " + notes_val if notes_val.strip() else ""
        else:
            notes_text = ""

        scalar_map.update({
            "CorrelationCheck.method2": cc.get("method", "") or "",
            "CorrelationCheck.threshold2": _fmt_float(cc.get("threshold")),
            "CorrelationCheck.plotted_features2": _fmt_int(cc.get("plotted_features")),
            "CorrelationCheck.n_numeric_features2": _fmt_int(cc.get("n_numeric_features")),
            "CorrelationCheck.n_pairs_flagged_ge_threshold2": _fmt_int(cc.get("n_pairs_flagged_ge_threshold")),
            "CorrelationCheck.n_pairs_total2": _fmt_int(cc.get("n_pairs_total")),
            "CorrelationCheck.top_pairs_csv_path": cc.get("top_pairs_csv") or cc.get("top_pairs_main_csv") or "",
            "CorrelationCheck.notes_text": notes_text,
        })
        scalar_map.update({
            "correlation_pearson_path": ctx.get("correlation_pearson_path", "") or "",
            "correlation_spearman_path": ctx.get("correlation_spearman_path", "") or "",
        })

        # VIF fallback text
        vif_ctx = ctx.get("VIFCheck") or {}
        vif_rows = vif_ctx.get("vif_table") or []
        scalar_map["VIFCheck.note_text"] = "" if (isinstance(vif_rows, list) and len(vif_rows) > 0) else "_No VIF results were generated for this run._"

        # Stress fallback text
        st_rows = (ctx.get("StressTestCheck") or {}).get("table") or []
        scalar_map["StressTest.note_text"] = "" if (isinstance(st_rows, list) and len(st_rows) > 0) else "_No stress-test results were generated for this run._"

        # InputCluster fallback text
        icc_rows = ctx.get("InputClusterCheck", {}).get("cluster_table") or []
        scalar_map["InputClusterCheck.note_text"] = "" if (isinstance(icc_rows, list) and len(icc_rows) > 0) else "_No cluster summary was generated for this run._"
        scalar_map["InputClusterCheck.cluster_csv"] = ctx.get("InputClusterCheck", {}).get("cluster_csv") or "(not available)"

        # RuleEngineCheck pretty text
        rec = ctx.get("RuleEngineCheck") or {}
        rules = rec.get("rules") or {}
        def _fmt_rule_value(v):
            if isinstance(v, bool):
                return "✅ Pass" if v else "❌ Fail"
            return str(v)
        overall_pass_val = rec.get("overall_pass")
        if isinstance(overall_pass_val, bool):
            overall_pass_str = "✅ Yes" if overall_pass_val else "❌ No"
        else:
            overall_pass_str = str(overall_pass_val) if overall_pass_val is not None else "—"
        if isinstance(rules, dict) and rules:
            lines = [f"- {str(k)}: {_fmt_rule_value(v)}" for k, v in rules.items()]
            rules_text = "\n".join(lines)
        else:
            rules_text = "No rule results were generated for this run."
        scalar_map.update({
            "RuleEngineCheck.overall_pass_str": overall_pass_str,
            "RuleEngineCheck.rules_text": rules_text,
        })

        # ModelMeta hyperparams & attributes text
        mm = ctx.get("ModelMetaCheck") or {}
        hyper_table = mm.get("hyperparam_table") or []
        if isinstance(hyper_table, list) and hyper_table:
            hp_lines = []
            for row in hyper_table:
                try:
                    k = row.get("param"); v = row.get("value")
                except Exception:
                    k, v = None, None
                if k is not None:
                    hp_lines.append(f"- {k}: {v}")
            hyperparams_text = "\n".join(hp_lines) if hp_lines else "_No hyperparameters provided._"
        else:
            hyperparams_text = "_No hyperparameters provided._"

        attrs = mm.get("attributes")
        if isinstance(attrs, dict) and attrs:
            attr_lines = [f"- {k}: {v}" for k, v in attrs.items()]
            attributes_text = "\n".join(attr_lines)
        else:
            attributes_text = (str(attrs) if attrs not in (None, "", []) else "_No attributes available._")

        scalar_map.update({
            "ModelMetaCheck.hyperparams_text": hyperparams_text,
            "ModelMetaCheck.attributes_text": attributes_text,
        })

        # Logistic stats (optional)
        lsf = ctx.get("LogisticStatsFit") or {}
        scalar_map.update({
            "LogisticStatsFit.log_lik2":   _fmt2(lsf.get("log_lik")),
            "LogisticStatsFit.aic2":       _fmt2(lsf.get("aic")),
            "LogisticStatsFit.bic2":       _fmt2(lsf.get("bic")),
            "LogisticStatsFit.pseudo_r22": _fmt2(lsf.get("pseudo_r2")),
            "LogitStats.summary_text": (
                ctx.get("LogisticStatsSummary_text")
                or ctx.get("LogisticStatsSummary")
                or "Logistic diagnostics not available."
            ),
        })

        # Linear (OLS) stats (optional)
        lin = ctx.get("LinearStats") or {}
        scalar_map["LinearStats.summary_text"] = lin.get("summary_text", "") or ""

        # EDA inclusion controls
        all_eda = ctx.get("eda_images_paths") or []
        opts = ctx.get("report_options") or {}
        K = opts.get("eda_max_images")
        if K is None:
            K = len(all_eda)
        else:
            try:
                K = int(K)
            except Exception:
                K = len(all_eda)
        eda_cols = int(opts.get("eda_grid_cols", 3))
        eda_subset = all_eda[:K] if K >= 0 else all_eda
        scalar_map["eda_count_note"] = f"(showing {len(eda_subset)} of {len(all_eda)})" if len(all_eda) != len(eda_subset) else ""
        ctx["_eda_subset"] = eda_subset
        ctx["_eda_cols"] = eda_cols

        # text replacements
        # --- CV Plot Generation (Spaghetti Plots) ---
        cv_curve_data = ctx.get("summary", {}).get("cv_stats", {}).get("curves")
        cv_plot_paths = {}

        if cv_curve_data:
            from pathlib import Path
            if "roc" in cv_curve_data and cv_curve_data["roc"]:
                out_p = TMP_DIR / "cv_roc_curve.png"
                _plot_cv_spaghetti(
                    cv_curve_data["roc"], 
                    title="ROC Curve (Cross-Validation)", 
                    xlabel="False Positive Rate", 
                    ylabel="True Positive Rate", 
                    output_path=str(out_p)
                )
                cv_plot_paths["cv_roc"] = str(out_p)

            if "pr" in cv_curve_data and cv_curve_data["pr"]:
                out_p = TMP_DIR / "cv_pr_curve.png"
                _plot_cv_spaghetti(
                    cv_curve_data["pr"], 
                    title="Precision-Recall Curve (Cross-Validation)", 
                    xlabel="Recall", 
                    ylabel="Precision", 
                    output_path=str(out_p)
                )
                cv_plot_paths["cv_pr"] = str(out_p)

        # --- OOF Lift/Gain/Regression ---
        oof_data = ctx.get("summary", {}).get("cv_stats", {}).get("oof")
        if oof_data and "y_true" in oof_data:
            # --- Classification CV Plots (requires y_prob) ---
            if "y_prob" in oof_data:
                # Lift
                out_p = TMP_DIR / "cv_lift_curve.png"
                _plot_oof_lift(
                    oof_data["y_true"],
                    oof_data["y_prob"],
                    output_path=str(out_p)
                )
                cv_plot_paths["cv_lift"] = str(out_p)
                
                # Calibration
                out_p2 = TMP_DIR / "cv_calib_curve.png"
                _plot_oof_calibration(
                    oof_data["y_true"],
                    oof_data["y_prob"],
                    output_path=str(out_p2)
                )
                cv_plot_paths["cv_calib"] = str(out_p2)

                # Confusion Matrix
                out_p_cm = TMP_DIR / "cv_confusion_matrix.png"
                _plot_cv_confusion_matrix(
                    oof_data["y_true"],
                    oof_data["y_pred"],
                    output_path=str(out_p_cm)
                )
                cv_plot_paths["cv_confusion"] = str(out_p_cm)

                # KS Curve
                out_p_ks = TMP_DIR / "cv_ks_curve.png"
                _plot_oof_ks(
                    oof_data["y_true"],
                    oof_data["y_prob"],
                    output_path=str(out_p_ks)
                )
                cv_plot_paths["cv_ks"] = str(out_p_ks)

                # Recompute Decile Lift Table
                try:
                    from plot_helper_table import build_lift_table_from_oof
                    cv_lift_rows = build_lift_table_from_oof(
                        oof_data["y_true"],
                        oof_data["y_prob"]
                    )
                    if "classification_tables" not in ctx: ctx["classification_tables"] = {}
                    ctx["classification_tables"]["lift"] = cv_lift_rows
                except Exception as e:
                    print(f"Error calculating CV lift table: {e}")

            # --- Regression CV Plots (y_pred only, no y_prob) ---
            elif "y_pred" in oof_data:
                # Import plotting helpers (lazy import inside method/if-block)
                try:
                    from plot_helper_reg import (
                         plot_reg_residuals_vs_pred,
                         plot_reg_residual_hist,
                         plot_reg_qq,
                         plot_reg_abs_error_box,
                         plot_reg_abs_error_violin
                    )
                    
                    # 1. Residuals vs Predicted
                    out_p1 = TMP_DIR / "cv_reg_residuals_vs_pred.png"
                    plot_reg_residuals_vs_pred(oof_data["y_true"], oof_data["y_pred"], str(out_p1))
                    cv_plot_paths["cv_reg_residuals_vs_pred"] = str(out_p1)
                    
                    # 2. Residual Histogram
                    out_p2 = TMP_DIR / "cv_reg_residual_hist.png"
                    plot_reg_residual_hist(oof_data["y_true"], oof_data["y_pred"], str(out_p2))
                    cv_plot_paths["cv_reg_residual_hist"] = str(out_p2)

                    # 3. Q-Q Plot
                    out_p3 = TMP_DIR / "cv_reg_qq.png"
                    plot_reg_qq(oof_data["y_true"], oof_data["y_pred"], str(out_p3))
                    cv_plot_paths["cv_reg_qq"] = str(out_p3)

                    # 4. Box Plot
                    out_p4 = TMP_DIR / "cv_reg_abs_error_box.png"
                    plot_reg_abs_error_box(oof_data["y_true"], oof_data["y_pred"], str(out_p4))
                    cv_plot_paths["cv_reg_abs_error_box"] = str(out_p4)

                    # 5. Violin Plot
                    out_p5 = TMP_DIR / "cv_reg_abs_error_violin.png"
                    plot_reg_abs_error_violin(oof_data["y_true"], oof_data["y_pred"], str(out_p5))
                    cv_plot_paths["cv_reg_abs_error_violin"] = str(out_p5)
                    
                    # 6. Pred vs Actual (Previously implemented)
                    out_p_reg = TMP_DIR / "cv_pred_vs_actual.png"
                    _plot_oof_pred_vs_actual(
                        oof_data["y_true"],
                        oof_data["y_pred"],
                        output_path=str(out_p_reg)
                    )
                    cv_plot_paths["cv_pred_vs_actual"] = str(out_p_reg)

                except Exception as e:
                    print(f"Failed to plot CV Regression Diagnostics: {e}")

        # Expose CV plot paths to Jinja2 context for conditional rendering
        scalar_map.update(cv_plot_paths)

        _replace_text_placeholders(doc, scalar_map)

        # image markers
        images_map = {
            "roc": ctx["classification_plot_paths"].get("roc"),
            "cv_roc": cv_plot_paths.get("cv_roc"),
            "pr": ctx["classification_plot_paths"].get("pr"),
            "cv_pr": cv_plot_paths.get("cv_pr"),
            "lift": ctx["classification_plot_paths"].get("lift"),
            "cv_lift": cv_plot_paths.get("cv_lift"),
            "calibration": ctx["classification_plot_paths"].get("calibration"),
            "cv_calibration": cv_plot_paths.get("cv_calibration"),
            "confusion": ctx["classification_plot_paths"].get("confusion"),
            "cv_confusion": cv_plot_paths.get("cv_confusion"),
            "ks": ctx.get("ks_curve_path"),
            "cv_ks": cv_plot_paths.get("cv_ks"),
            # Regression CV additions
            "cv_reg_pred_vs_actual": cv_plot_paths.get("cv_pred_vs_actual"),
            "cv_reg_residuals_vs_pred": cv_plot_paths.get("cv_reg_residuals_vs_pred"),
            "cv_reg_residual_hist": cv_plot_paths.get("cv_reg_residual_hist"),
            "cv_reg_qq": cv_plot_paths.get("cv_reg_qq"),
            "cv_reg_abs_error_box": cv_plot_paths.get("cv_reg_abs_error_box"),
            "cv_reg_abs_error_violin": cv_plot_paths.get("cv_reg_abs_error_violin"),
            
            "correlation_heatmap": ctx.get("correlation_heatmap_path"),
            "shap_beeswarm": ctx.get("shap_beeswarm_path"),
            "shap_bar": ctx.get("shap_bar_path"),
            "reg_pred_vs_actual": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("pred_vs_actual"),
            "reg_residuals_vs_pred": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("residuals_vs_pred"),
            "reg_residual_hist": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("residual_hist"),
            "reg_qq": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("qq_plot"),
            "reg_abs_error_box": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("abs_error_box"),
            "reg_abs_error_violin": ctx.get("RegressionMetrics", {}).get("artifacts", {}).get("abs_error_violin"),
            "cluster_plot": plot_path,
            "logit_summary": ctx.get("logit_summary_path"),
        }
        _replace_image_markers(doc, images_map, width_in=4.8)

        # save once before inserting tables/grids by anchors
        doc.save(self.output_path)

        # ===== 3) Post-render insertions via anchors ==========================
        docx = Document(self.output_path)

        # EDA grid
        insert_image_grid(
            docx,
            anchor="Distribution Plots",
            img_paths=ctx.get("_eda_subset") or [],
            cols=ctx.get("_eda_cols") or 3,
            width_in=2.2
        )

        # helpers
        def _round_row_numbers(row, places=2):
            if not isinstance(row, dict):
                return row
            out = {}
            for k, v in row.items():
                try:
                    if v is None:
                        out[k] = v
                    else:
                        fv = float(v)
                        out[k] = round(fv, places) if np.isfinite(fv) else v
                except Exception:
                    out[k] = v
            return out

        stress_rows = (ctx.get("StressTestCheck", {}) or {}).get("table", []) or []
        stress_rows = [_round_row_numbers(r) for r in stress_rows]

        # OLS coefficients (rounded)
        ols_coeff_rows = (ctx.get("LinearStats") or {}).get("coeff_table", []) or []
        ols_coeff_rows = [_round_row_numbers(r, places=4) for r in ols_coeff_rows]

        # Ensure logistic minimal coef table exists/rounded (legacy 2-col)
        if "coef_table" not in ctx or not isinstance(ctx["coef_table"], list):
            ctx["coef_table"] = []
        logit_coeff_rows = ctx.get("coef_table", []) or []
        logit_coeff_rows = [_round_row_numbers(r, places=4) for r in logit_coeff_rows]
        ctx["coef_table"] = logit_coeff_rows

        # ---------- Build full logistic coefficients table (OLS-like) ----------
        def _coerce_float_or_blank(v):
            try:
                if v is None:
                    return ""
                fv = float(v)
                return fv if np.isfinite(fv) else ""
            except Exception:
                return ""

        logit_full_rows = []

        
        stats_full = None
        for source_key in ("LogisticStats", "LogisticStatsCheck"):
            cand = self._grab(source_key) or {}
            if isinstance(cand, dict) and isinstance(cand.get("coef_table_full"), list):
                stats_full = cand["coef_table_full"]
                break
        if isinstance(stats_full, list) and stats_full:
            for r in stats_full:
                logit_full_rows.append({
                    "feature":    str(r.get("feature")),
                    "coef":       _coerce_float_or_blank(r.get("coef")),
                    "std err":    _coerce_float_or_blank(r.get("std_err")),
                    "z":          _coerce_float_or_blank(r.get("z")),
                    "P>|z|":      _coerce_float_or_blank(r.get("p") or r.get("p_value") or r.get("p>|z|")),
                    "ci_low":     _coerce_float_or_blank(r.get("ci_low")),
                    "ci_high":    _coerce_float_or_blank(r.get("ci_high")),
                })

        # (2) fallback: synthesize from sklearn-like attributes (coef_ / intercept_)
        if not logit_full_rows and (ctx.get("task_type") or "").lower() == "classification":
            mm = ctx.get("ModelMetaCheck") or {}
            attrs = (mm.get("attributes") or {})
            feat_names = mm.get("feature_names") or []
            def _flatten(x):
                if x is None: return None
                arr = np.array(x)
                return arr.flatten().tolist()
            flat_coefs = _flatten(attrs.get("coef_"))
            intercept = attrs.get("intercept_")
            if isinstance(flat_coefs, list) and len(flat_coefs) == len(feat_names):
                for f, c in zip(feat_names, flat_coefs):
                    logit_full_rows.append({
                        "feature": str(f),
                        "coef": _coerce_float_or_blank(c),
                        "std err": "", "z": "", "P>|z|": "", "ci_low": "", "ci_high": ""
                    })
                if intercept is not None:
                    try:
                        b0 = float(intercept[0] if isinstance(intercept, (list, tuple, np.ndarray)) else intercept)
                    except Exception:
                        b0 = None
                    logit_full_rows.insert(0, {
                        "feature": "Intercept",
                        "coef": _coerce_float_or_blank(b0),
                        "std err": "", "z": "", "P>|z|": "", "ci_low": "", "ci_high": ""
                    })

        ctx["logit_coef_full_rows"] = logit_full_rows

        # ---------- task-aware table list ----------
        task = (ctx.get("task_type") or "classification").lower()

        if task == "regression":
            stress_headers = ["feature", "perturbation", "rmse", "r2", "delta_rmse", "delta_r2"]
            tbl_specs = [
                {
                    "anchor": "Top High-Correlation Feature Pairs (|r| ≥ Threshold)",
                    "headers": ["feature_i", "feature_j", "corr", "n_used", "pct_missing_i", "pct_missing_j"],
                    "rows": corr_preview_rows,
                },
                {
                    "anchor": "Stress Testing Results",
                    "headers": stress_headers,
                    "rows": stress_rows,
                },
                {
                    "anchor": "Cluster Summary Table",
                    "headers": ["Cluster", "Count", "Percent"],
                    "rows": ctx.get("InputClusterCheck", {}).get("cluster_table", []),
                },
                {
                    "anchor": "Variance Inflation Factor (VIF) Check",
                    "headers": ["Feature", "VIF"],
                    "rows": ctx.get("VIFCheck", {}).get("vif_table", []),
                },
                {
                    "anchor": "OLS Coefficients (Regression)",
                    "headers": ["feature", "coef", "std err", "t", "P>|t|", "ci_low", "ci_high"],
                    "rows": ols_coeff_rows,
                },
                {
                    "anchor": "Top SHAP Features",
                    "headers": ["feature", "mean_abs_shap"],
                    "rows": ctx.get("shap_top_features", []) or [],
                },
            ]
        else:
            stress_headers = ["feature", "perturbation", "accuracy", "auc", "delta_accuracy", "delta_auc"]

            labels = None
            try:
                labels = (cls.get("labels") or cls_tables.get("labels") or ctx.get("classification_labels"))
            except Exception:
                labels = None

            def _mk_confusion_headers(labs):
                if not labs or not isinstance(labs, (list, tuple)) or len(labs) == 0:
                    return ["", "Pred 0", "Pred 1"]
                return [""] + [f"Pred {str(l)}" for l in labs]

            confusion_rows = ctx.get("classification_tables", {}).get("confusion", []) or []
            confusion_headers = _mk_confusion_headers(labels)

            logit_stats = ctx.get("LogitStats") or {}
            logit_headers = logit_stats.get("coef_table_headers") or ["feature","coef","std err","z","P>|z|","ci_low","ci_high"]
            logit_rows = (
                logit_stats.get("coef_table_rows")
                or ctx.get("logit_coef_full_rows")
                or ctx.get("coef_table")
                or []
            )

            baseline_logit_rows = self._extract_baseline_logit_rows(ctx)

            tbl_specs = [
                {
                    "anchor": "Top High-Correlation Feature Pairs (|r| ≥ Threshold)",
                    "headers": ["feature_i", "feature_j", "corr", "n_used", "pct_missing_i", "pct_missing_j"],
                    "rows": corr_preview_rows,
                },
                {
                    "anchor": "Stress Testing Results",
                    "headers": stress_headers,
                    "rows": stress_rows,
                },
                {
                    "anchor": "Cluster Summary Table",
                    "headers": ["Cluster", "Count", "Percent"],
                    "rows": ctx.get("InputClusterCheck", {}).get("cluster_table", []),
                },
                {
                    "anchor": "Variance Inflation Factor (VIF) Check",
                    "headers": ["Feature", "VIF"],
                    "rows": ctx.get("VIFCheck", {}).get("vif_table", []),
                },
                {
                    "anchor": "Confusion Matrix (Classification)",
                    "headers": confusion_headers,
                    "rows": confusion_rows if confusion_rows else [{"": "(no confusion matrix available)"}],
                },
                {
                    "anchor": "Decile Lift Table",
                    "headers": ["decile", "total", "events", "avg_score", "event_rate", "lift",
                                "cum_events", "cum_total", "cum_capture_rate", "cum_population", "cum_gain"],
                    "rows": ctx.get("classification_tables", {}).get("lift", []),
                },
                {
                    "anchor": "Baseline Metrics",
                    "headers": ["metric", "value"],
                    "rows": baseline_logit_rows,
                },
                {
                    "anchor": "Logistic Regression Coefficients ",
                    "headers": logit_headers,
                    "rows": logit_rows if logit_rows else [
                        {"feature": "(no coefficients available)", "coef": "", "std err": "", "z": "", "P>|z|": "", "ci_low": "", "ci_high": ""}
                    ],
                },
                {
                    "anchor": "Feature Importances (Tree-Based Models)",
                    "headers": ["feature", "importance"],
                    "rows": ctx.get("feature_importance_table", []) or [],
                },
                {
                    "anchor": "Top SHAP Features",
                    "headers": ["feature", "mean_abs_shap"],
                    "rows": ctx.get("shap_top_features", []) or [],
                },
            ]

        for spec in tbl_specs:
            if spec["rows"]:
                if spec["anchor"] == "Decile Lift Table":
                    # --- Render as Image ---
                    img_path = TMP_DIR / "lift_table_img.png"
                    _render_table_as_image(spec["headers"], spec["rows"], str(img_path))
                    insert_image_after(docx, spec["anchor"], str(img_path), width_in=6.5)
                    print(f"✅ added table IMAGE after «{spec['anchor']}»")
                else:
                    tbl = build_table(docx, spec["headers"], spec["rows"])
                    insert_after(docx, spec["anchor"], tbl)
                    print(f"✅ added table after «{spec['anchor']}»")


        docx.save(self.output_path)

def _plot_cv_spaghetti(curve_data, title, xlabel, ylabel, output_path):
    """
    Generates a spaghetti plot for CV curves (ROC or PR).
    curve_data: list of (x, y) tuples, one per fold.
    output_path: where to save the image.
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    plt.figure(figsize=(6, 5))
    
    # Interpolate to common x-axis for mean curve calculation
    common_x = np.linspace(0, 1, 100)
    interp_ys = []
    
    # Plot individual folds
    for x, y in curve_data:
        plt.plot(x, y, color='gray', alpha=0.3, lw=1)
        # Interpolate
        # Note: np.interp needs sorted x. ROC x (FPR) is sorted. PR x (Recall) is sorted descending.
        # For PR curve (Recall on x), we need to handle sorting.
        if x[0] > x[-1]: # Descending
             interp_ys.append(np.interp(common_x, x[::-1], y[::-1]))
        else:
             interp_ys.append(np.interp(common_x, x, y))
        
    # Calculate and plot mean curve
    mean_y = np.mean(interp_ys, axis=0)
    std_y = np.std(interp_ys, axis=0)
    
    plt.plot(common_x, mean_y, color='C0', lw=2, label='Mean CV')
    
    # Add +/- 1 std deviation shading
    plt.fill_between(common_x, 
                     np.maximum(mean_y - std_y, 0), 
                     np.minimum(mean_y + std_y, 1), 
                     color='C0', alpha=0.2, label=r'$\pm$ 1 std. dev.')
    
    # Formatting
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.legend(loc="lower right" if "Recall" not in xlabel else "lower left")
    plt.grid(alpha=0.3)
    plt.xlim([-0.01, 1.01])
    plt.ylim([-0.01, 1.01])
    
    # Specific diagonal line for ROC
    if "False Positive" in xlabel:
         plt.plot([0, 1], [0, 1], linestyle='--', lw=1, color='r', alpha=0.8, label='Chance')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()

def _plot_oof_lift(y_true, y_prob, output_path):
    """
    Plots Cumulative Gain / Lift curve from OOF predictions.
    Computes simple lift: (Precision at K) / (Global Pct of Positives)
    """
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    
    # Create DF and sort by probability descending
    df = pd.DataFrame({"y": y_true, "p": y_prob})
    df = df.sort_values("p", ascending=False).reset_index(drop=True)
    
    total_pos = df["y"].sum()
    n = len(df)
    if total_pos == 0: return # robust
    
    # Cumulative positives
    df["cum_pos"] = df["y"].cumsum()
    # Pct of positives captured
    df["gain"] = df["cum_pos"] / total_pos
    # Pct of population
    df["pop_pct"] = (df.index + 1) / n
    
    # Lift = Gain / Pop_Pct
    # But often "Lift Curve" plots Lift vs Pop_Pct. "Gain Curve" plots Gain vs Pop_Pct.
    # The existing template usually shows "Cumulative Gain / Lift" which is often the Gain chart compared to random.
    # Let's verify existing plot style. Usually it's Gain (Captured Response) vs Population.
    # Random line is y=x.
    
    plt.figure(figsize=(6, 5))
    plt.plot(df["pop_pct"], df["gain"], label="Model (OOF)", color="C0", lw=2)
    plt.plot([0, 1], [0, 1], 'r--', label="Random Model")
    
    # Optional: "Wizard" / Perfect Model
    # Sort perfect would be all 1s then all 0s
    perfect_k = int(total_pos)
    perfect_gain = np.concatenate([
        np.linspace(0, 1, perfect_k),
        np.ones(n - perfect_k)
    ])
    # This is rough approximation for plotting 'Perfect' if desired, but let's stick to standard Gain
    
    plt.xlabel("% of Population Contacted")
    plt.ylabel("% of Targets Captured (Cumulative Gain)")
    plt.title("Cumulative Gain (OOF Pooled)")
    plt.legend()
    plt.grid(alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()

def _plot_oof_calibration(y_true, y_prob, output_path):
    """
    Plots Calibration (Reliability) curve from OOF predictions.
    """
    import matplotlib.pyplot as plt
    from sklearn.calibration import calibration_curve
    
    # Calculate calibration
    prob_true, prob_pred = calibration_curve(y_true, y_prob, n_bins=10, strategy='uniform')
    
    plt.figure(figsize=(6, 5))
    
    # Perfect calibration line
    plt.plot([0, 1], [0, 1], linestyle='--', color='gray', label='Perfectly Calibrated')
    
    # Model calibration
    plt.plot(prob_pred, prob_true, marker='o', linewidth=1, label='Model (OOF)', color='C0')
    
    plt.xlabel("Mean Predicted Probability")
    plt.ylabel("Fraction of Positives")
    plt.title("Calibration Curve (Reliability Diagram)")
    plt.legend()
    plt.grid(alpha=0.3)
    plt.xlim([-0.05, 1.05])
    plt.ylim([-0.05, 1.05])
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()

def _plot_cv_confusion_matrix(y_true, y_pred, output_path):
    """
    Plots a Confusion Matrix from OOF predictions.
    """
    import matplotlib.pyplot as plt
    import seaborn as sns
    from sklearn.metrics import confusion_matrix
    import numpy as np

    cm = confusion_matrix(y_true, y_pred)
    
    # Normalize for color mapping, but annotate with raw counts
    cm_norm = cm.astype('float') / cm.sum(axis=1)[:, np.newaxis]
    
    plt.figure(figsize=(6, 5))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', cbar=False)
    
    plt.xlabel('Predicted Label')
    plt.ylabel('True Label')
    plt.title('Confusion Matrix (Pooled CV)')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()

def _calculate_oof_lift_table(y_true, y_prob):
    """
    Calculates Decile Lift Table rows from OOF predictions.
    Returns a list of lists matching the headers:
    ["decile", "total", "events", "avg_score", "event_rate", "lift",
     "cum_events", "cum_total", "cum_capture_rate", "cum_population", "cum_gain"]
    
    cum_gain here is interpreted as Cumulative Lift.
    """
    import pandas as pd
    import numpy as np

    df = pd.DataFrame({"y_true": y_true, "y_prob": y_prob})
    df = df.sort_values("y_prob", ascending=False).reset_index(drop=True)
    
    # Global stats
    total_events = df["y_true"].sum()
    total_pop = len(df)
    global_rate = total_events / total_pop if total_pop > 0 else 0
    
    # Deciles (1 to 10)
    # properly handle cases with fewer than 10 rows
    n_bins = 10
    df["decile"] = pd.qcut(df.index, q=n_bins, labels=False, duplicates='drop') + 1
    # qcut does equal size buckets. But since we sorted desc, decile 1 is top.
    # Actually qcut on index (0..N) with sorted DF:
    # index 0 is top score. 
    # pd.qcut on default creates bins based on values.
    # We want equal SIZE buckets.
    # simplest:
    df["decile"] = np.ceil((df.index + 1) / len(df) * 10).astype(int)
    
    # Aggregation
    g = df.groupby("decile")
    agg = g.agg({
        "y_true": ["count", "sum"],
        "y_prob": "mean"
    })
    
    # Flatten cols
    # agg.columns is (y_true, count), (y_true, sum), (y_prob, mean)
    counts = agg[("y_true", "count")]
    events = agg[("y_true", "sum")]
    avg_score = agg[("y_prob", "mean")]
    
    rows = []
    
    # Cumulative stats
    cum_events = 0
    cum_total = 0
    
    for d in range(1, 11):
        if d not in counts.index:
            continue
            
        n = counts.loc[d]
        e = events.loc[d]
        score = avg_score.loc[d]
        
        rate = e / n if n > 0 else 0
        lift = rate / global_rate if global_rate > 0 else 0
        
        cum_events += e
        cum_total += n
        
        cum_capture = cum_events / total_events if total_events > 0 else 0
        cum_pop = cum_total / total_pop if total_pop > 0 else 0
        
        cum_rate = cum_events / cum_total if cum_total > 0 else 0
        cum_lift_val = cum_rate / global_rate if global_rate > 0 else 0
        
        # Row format: 
        # ["decile", "total", "events", "avg_score", "event_rate", "lift",
        #  "cum_events", "cum_total", "cum_capture_rate", "cum_population", "cum_gain"]
        
        row = [
            d,              # decile
            int(n),         # total
            int(e),         # events
            score,          # avg_score
            rate,           # event_rate
            lift,           # lift
            int(cum_events),# cum_events
            int(cum_total), # cum_total
            cum_capture,    # cum_capture_rate
            cum_pop,        # cum_population
            cum_lift_val    # cum_gain (Cum Lift)
        ]
        rows.append(row)
        
    return rows

def _render_table_as_image(headers, rows, output_path):
    """
    Renders a table as an image using matplotlib.
    Handles formatting (rounding) automatically.
    """
    import matplotlib.pyplot as plt
    import pandas as pd
    import numpy as np

    # Create DF
    df = pd.DataFrame(rows, columns=headers)
    
    # Format Mapping
    # Identify int columns and float columns
    # We know the headers for Lift Table:
    # ["decile", "total", "events", "avg_score", "event_rate", "lift", 
    #  "cum_events", "cum_total", "cum_capture_rate", "cum_population", "cum_gain"]
    
    format_dict = {}
    
    # Ints
    for col in ["decile", "total", "events", "cum_events", "cum_total"]:
        if col in df.columns:
            # fillna(0) and convert to int for safety, then string with comma
            df[col] = df[col].fillna(0).astype(int).apply(lambda x: f"{x:,}")
            
    # Floats (4 decimals)
    for col in ["avg_score", "event_rate", "cum_capture_rate", "cum_population"]:
        if col in df.columns:
            df[col] = df[col].astype(float).apply(lambda x: f"{x:.4f}")
            
    # Floats (2 decimals for Lift)
    for col in ["lift", "cum_gain"]:
        if col in df.columns:
            df[col] = df[col].astype(float).apply(lambda x: f"{x:.2f}")

    # Clean Headers (User friendly)
    clean_map = {
        "decile": "Decile", "total": "Total", "events": "Events",
        "avg_score": "Avg Score", "event_rate": "Event Rate", "lift": "Lift",
        "cum_events": "Cum Events", "cum_total": "Cum Total",
        "cum_capture_rate": "Cum Capture", "cum_population": "Cum Pop",
        "cum_gain": "Cum Lift"
    }
    pretty_cols = [clean_map.get(c, c) for c in df.columns]

    # Plot
    # Aspect ratio: wide
    h, w = df.shape
    fig_height = h * 0.4 + 0.8
    fig_width = w * 1.2
    
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('tight')
    ax.axis('off')
    
    # Table
    table = ax.table(cellText=df.values,
                     colLabels=pretty_cols,
                     cellLoc='center',
                     loc='center')
    
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1.2, 1.5) # x_scale, y_scale
    
    # Bold Headers
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#f0f0f0')
            
    plt.tight_layout()
    plt.savefig(output_path, dpi=200, bbox_inches='tight', pad_inches=0.1)
    plt.close()

def _plot_oof_ks(y_true, y_prob, output_path):
    """
    Plots the KS Curve (CDF of Events vs Non-Events) for Pooled OOF predictions.
    Replicates the visual style of the standard KS plot.
    """
    import matplotlib.pyplot as plt
    import pandas as pd
    import numpy as np
    
    # 1. Prepare Dataframe
    # We can infer pos_label=1 usually
    pos_label = 1
    df = pd.DataFrame({"y": (np.array(y_true) == pos_label).astype(int), "score": np.array(y_prob)})
    
    if df.empty:
        return # Handle empty

    df = df.sort_values("score", ascending=False).reset_index(drop=True)
    n = len(df)

    # counts
    total_events = df["y"].sum()
    total_non_events = n - total_events

    # avoid divide-by-zero
    if total_events == 0 or total_non_events == 0:
        return

    cum_events = np.cumsum(df["y"].values) / total_events
    cum_non_events = np.cumsum(1 - df["y"].values) / total_non_events
    population = (np.arange(1, n + 1)) / n
    ks_gap = np.abs(cum_events - cum_non_events)
    
    ks_df = pd.DataFrame({
        "population": population,
        "cum_event": cum_events,
        "cum_non_event": cum_non_events,
        "ks_gap": ks_gap
    })
    
    # 2. Plotting
    ks_idx       = int(ks_df["ks_gap"].values.argmax())
    ks_x         = float(ks_df["population"].iloc[ks_idx])
    ks_y_event   = float(ks_df["cum_event"].iloc[ks_idx])
    ks_y_nonevent= float(ks_df["cum_non_event"].iloc[ks_idx])
    ks_val_annot = abs(ks_y_event - ks_y_nonevent)

    plt.figure(figsize=(6, 4))
    plt.plot(ks_df["population"], ks_df["cum_event"],     label="Cumulative Event")
    plt.plot(ks_df["population"], ks_df["cum_non_event"], label="Cumulative Non-Event")

    # vertical line & markers at max KS
    plt.axvline(ks_x, linestyle="--", alpha=0.7, color='gray')
    plt.scatter([ks_x], [ks_y_event],    s=25, color='C0')
    plt.scatter([ks_x], [ks_y_nonevent], s=25, color='C1')

    # readable annotation
    plt.annotate(
        f"KS = {ks_val_annot:.1%}\nPop = {ks_x:.1%}",
        xy=(ks_x, (ks_y_event + ks_y_nonevent) / 2.0),
        xytext=(ks_x + 0.05, min(0.9, (ks_y_event + ks_y_nonevent) / 2.0 + 0.1)),
        arrowprops=dict(arrowstyle="->", color="black", lw=1),
        ha="left", va="center", fontsize=10,
        bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="black", lw=0.8),
    )

    plt.xlabel("Population (fraction)")
    plt.ylabel("Cumulative share")
    plt.title("CV: KS Curve (Pooled OOF)")
    plt.legend(loc="lower right")
    plt.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=160)
    plt.close()

def _plot_oof_pred_vs_actual(y_true, y_pred, output_path):
    """
    Plots Predicted vs Actual for Pooled OOF regression predictions.
    Replicates the visual style of the standard regression plot.
    """
    import matplotlib.pyplot as plt
    import numpy as np

    plt.figure()
    plt.scatter(y_true, y_pred, s=12, alpha=0.5, label="OOF Predictions")
    
    # Check for empty or single-point data to avoid min/max errors
    if len(y_true) > 0:
        mn = float(min(np.min(y_true), np.min(y_pred)))
        mx = float(max(np.max(y_true), np.max(y_pred)))
        plt.plot([mn, mx], [mn, mx], 'k--', lw=1, label="Perfect Fit")
    
    plt.xlabel("Actual")
    plt.ylabel("Predicted")
    plt.title("CV: Predicted vs Actual (Pooled OOF)")
    plt.legend(loc="upper left")
    plt.tight_layout()
    plt.savefig(output_path, dpi=160)
    plt.close()
