from docx import Document
from tanml.report.report_builder import _replace_placeholder_with_kv_table
import os

# Create a dummy doc
doc = Document()
doc.add_paragraph("Some text before.")
doc.add_paragraph("Meta: {{ModelMetaCheck.target_balance}}")
doc.add_paragraph("Some text after.")

# Create a table and put placeholder in cell
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]
cell.text = "In table: {{ModelMetaCheck.target_balance}}"

# Create a nested table test
table2 = doc.add_table(rows=1, cols=1)
cell2 = table2.rows[0].cells[0]
nested_table = cell2.add_table(rows=1, cols=1)
nested_cell = nested_table.rows[0].cells[0]
nested_cell.text = "In nested table: {{ModelMetaCheck.target_balance}}"

# Create a SPACED placeholder test
table3 = doc.add_table(rows=1, cols=1)
cell3 = table3.rows[0].cells[0]
cell3.text = "Spaced test: {{ ModelMetaCheck.target_balance }}"

data = {"Mean": "2.40", "Std": "1.00"}

print("Testing body replacement...")
res1 = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", data)
print(f"Body Replacement Success: {res1}")

print("Testing table replacement...")
res2 = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", data)
print(f"Table Replacement Success: {res2}")

print("Testing nested table replacement...")
res3 = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", data)
print(f"Nested Table Replacement Success: {res3}")

print("Testing SPACED nested table replacement...")
res4 = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", data)
print(f"Spaced Table Replacement Success: {res4}")

# Create SPACED BODY test (since real failure was in body)
doc.add_paragraph("Spaced Body: {{ ModelMetaCheck.target_balance }}")
print("Testing SPACED BODY replacement...")
res5 = _replace_placeholder_with_kv_table(doc, "{{ModelMetaCheck.target_balance}}", data)
print(f"Spaced Body Replacement Success: {res5}")


doc.save("test_table_replace.docx")
print("Saved test_table_replace.docx")
